from pathlib import Path

import pytest
from docx import Document

from services.firma import plantillas_service as module
from services.firma.plantillas_service import PlantillasFirmaService


def _docx(path: Path):
    doc = Document()
    doc.add_paragraph("Autorizacion de {{cliente_nombre}} con NIF {{cliente_nif}}")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Fecha: {{fecha_documento}}"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Expediente {{expediente}}"
    doc.save(path)


def _patch_dirs(monkeypatch, root: Path):
    monkeypatch.setattr(module, "get_word_templates_dir", lambda: str(root))

    def subdir(tipo, crear=True):
        path = root / tipo
        if crear:
            path.mkdir(parents=True, exist_ok=True)
        return path

    monkeypatch.setattr(module, "get_word_templates_subdir", subdir)


def test_detecta_campos_en_parrafos_tablas_y_cabeceras(tmp_path):
    path = tmp_path / "modelo.docx"
    _docx(path)

    assert PlantillasFirmaService.detectar_campos(path) == [
        "cliente_nombre", "cliente_nif", "fecha_documento", "expediente"
    ]


def test_organizador_copia_sin_borrar_y_detecta_conflictos(monkeypatch, tmp_path):
    _patch_dirs(monkeypatch, tmp_path)
    source = tmp_path / "factura.docx"
    source.write_bytes(b"word-1")

    first = PlantillasFirmaService.copiar_clasificadas({"factura.docx": "facturas"})
    second = PlantillasFirmaService.copiar_clasificadas({"factura.docx": "facturas"})
    (tmp_path / "facturas" / "factura.docx").write_bytes(b"otro")
    third = PlantillasFirmaService.copiar_clasificadas({"factura.docx": "facturas"})

    assert first["copiados"] == ["factura.docx"]
    assert second["identicos"] == ["factura.docx"]
    assert third["conflictos"] == ["factura.docx"]
    assert source.read_bytes() == b"word-1"


def test_precarga_por_origen_sin_modificar_maestros():
    service = PlantillasFirmaService(object())
    plantilla = {"campos": [
        {"clave": "nombre", "origen": "empresa", "campo_origen": "nombre", "valor_defecto": ""},
        {"clave": "correo", "origen": "tercero", "campo_origen": "email", "valor_defecto": ""},
        {"clave": "nota", "origen": "manual", "valor_defecto": "Introducir"},
    ]}
    empresa = {"nombre": "Cliente SL"}
    tercero = {"email": "persona@example.com"}

    values = service.valores_iniciales(plantilla, empresa, tercero, "Gestor")

    assert values == {"nombre": "Cliente SL", "correo": "persona@example.com", "nota": "Introducir"}
    assert empresa == {"nombre": "Cliente SL"}
    assert tercero == {"email": "persona@example.com"}


def test_valida_obligatorios_email_y_fecha():
    plantilla = {"campos": [
        {"clave": "email", "etiqueta": "Email", "tipo": "email", "obligatorio": 1},
        {"clave": "fecha", "etiqueta": "Fecha", "tipo": "fecha", "obligatorio": 1},
    ]}

    PlantillasFirmaService.validar_valores(
        plantilla, {"email": "cliente@example.com", "fecha": "08/08/2026"}
    )
    with pytest.raises(ValueError, match="email valido"):
        PlantillasFirmaService.validar_valores(
            plantilla, {"email": "incorrecto", "fecha": "08/08/2026"}
        )
