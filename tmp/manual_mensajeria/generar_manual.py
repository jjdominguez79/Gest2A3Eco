from __future__ import annotations

import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WORK = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "docs" / "Manual_Mensajeria_Gestinem.docx"
OUTPUT_COPY = ROOT / "output" / "documentos" / OUT_DOCX.name
LOGO = ROOT / "gestinem_app" / "assets" / "images" / "logo_new.png"
ICON = ROOT / "gestinem_app" / "assets" / "icons" / "app_icon.png"
URL = "https://app.gestinem.es"

NAVY = "0D2A6B"
BLUE = "0759AF"
LIGHT = "EEF4FA"
PALE = "F7F9FC"
INK = "183247"
MUTED = "5D6B78"
GREEN = "18794E"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def image_font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\seguisb.ttf") if bold else Path(r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf") if bold else Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_browser_guide() -> Path:
    canvas = Image.new("RGB", (1200, 760), "white")
    draw = ImageDraw.Draw(canvas)
    rounded(draw, (45, 35, 1155, 725), 28, "#F4F7FA", outline="#CFD9E2", width=3)
    rounded(draw, (45, 35, 1155, 135), 28, "#E6EBF0")
    for x, color in ((90, "#F16B61"), (130, "#F6C350"), (170, "#65C466")):
        draw.ellipse((x - 12, 72, x + 12, 96), fill=color)
    rounded(draw, (235, 60, 1075, 112), 24, "white", outline="#C7D2DC")
    draw.text((270, 86), URL, font=image_font(24), fill="#34495E", anchor="lm")

    icon = Image.open(ICON).convert("RGBA")
    icon.thumbnail((145, 145), Image.Resampling.LANCZOS)
    canvas.paste(icon, (160, 220), icon)
    draw.text((350, 230), "Bienvenido a Gestinem", font=image_font(38, True), fill="#0D2A6B")
    draw.text((350, 290), "Mensajería y documentos con tu despacho", font=image_font(24), fill="#5D6B78")
    rounded(draw, (350, 365, 995, 430), 12, "white", outline="#B8C6D2", width=2)
    draw.text((380, 398), "Correo electrónico", font=image_font(23), fill="#73808C", anchor="lm")
    rounded(draw, (350, 460, 995, 525), 12, "white", outline="#B8C6D2", width=2)
    draw.text((380, 493), "Contraseña", font=image_font(23), fill="#73808C", anchor="lm")
    rounded(draw, (350, 565, 995, 635), 12, "#0759AF")
    draw.text((672, 600), "Entrar", font=image_font(27, True), fill="white", anchor="mm")
    path = WORK / "01-acceso-web.png"
    canvas.save(path, optimize=True)
    return path


def make_channels_guide() -> Path:
    canvas = Image.new("RGB", (1000, 1180), "#F4F7FA")
    draw = ImageDraw.Draw(canvas)
    rounded(draw, (55, 35, 945, 1145), 32, "white", outline="#D6E0E8", width=3)
    draw.text((105, 105), "Conversaciones", font=image_font(38, True), fill="#0D2A6B")
    draw.text((105, 160), "Selecciona el tema antes de escribir", font=image_font(23), fill="#5D6B78")
    items = [
        ("LA", "Laboral", "Nóminas, contratos, altas y bajas"),
        ("CF", "Contable / Fiscal", "Facturas, impuestos y contabilidad"),
        ("DP", "Directo", "Conversación reservada con el despacho"),
    ]
    y = 245
    for initials, title, detail in items:
        rounded(draw, (95, y, 905, y + 190), 20, "#F7F9FC", outline="#DCE4EA", width=2)
        draw.ellipse((130, y + 48, 220, y + 138), fill="#0759AF")
        draw.text((175, y + 93), initials, font=image_font(25, True), fill="white", anchor="mm")
        draw.text((260, y + 60), title, font=image_font(29, True), fill="#0D2A6B")
        draw.text((260, y + 112), detail, font=image_font(22), fill="#5D6B78")
        draw.text((850, y + 95), ">", font=image_font(35, True), fill="#0759AF", anchor="mm")
        y += 220
    rounded(draw, (95, 940, 905, 1060), 18, "#EAF6EF", outline="#B9DFC8", width=2)
    draw.text((500, 980), "Consejo", font=image_font(24, True), fill="#18794E", anchor="mm")
    draw.text((500, 1025), "Usa el canal correcto para recibir una respuesta más rápida.", font=image_font(21), fill="#315F48", anchor="mm")
    path = WORK / "02-canales.png"
    canvas.save(path, optimize=True)
    return path


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_style(style, size, color, bold, before, after, line=1.12):
    style.font.name = "Aptos"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def shade(paragraph, fill):
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(element)


def add_callout(doc, title, body, fill=LIGHT, accent=NAVY):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.12
    shade(paragraph, fill)
    set_font(paragraph.add_run(title + "\n"), size=11, color=accent, bold=True)
    set_font(paragraph.add_run(body), size=10.5)


def configure_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    target = "bullet" if bullet else "decimal"
    namespace = {"w": qn("w:val").split("}")[0][1:]}
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        fmt = abstract.find("./w:lvl[@w:ilvl='0']/w:numFmt", namespace)
        if fmt is not None and fmt.get(qn("w:val")) == target:
            abstract_id = int(abstract.get(qn("w:abstractNumId")))
            break
    if abstract_id is None:
        raise RuntimeError(f"No existe una lista {target} en la plantilla")
    num_id = max(
        [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))],
        default=0,
    ) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    if not bullet:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        num.append(override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.19)
    paragraph.paragraph_format.space_after = Pt(5)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph._p.get_or_add_pPr().append(num_pr)
    set_font(paragraph.add_run(text))


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([props, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_figure(doc, path, width, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(8)
    set_font(label.add_run(caption), size=8.5, color=MUTED, italic=True)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, size=8.5, color=MUTED)


def new_page(doc, title, subtitle=""):
    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(title)
    if subtitle:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
        set_font(paragraph.add_run(subtitle), color=MUTED)


def build_document():
    browser_guide = make_browser_guide()
    channels_guide = make_channels_guide()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.footer_distance = Inches(0.42)

    set_style(doc.styles["Normal"], 11, INK, False, 0, 6, 1.18)
    set_style(doc.styles["Heading 1"], 17, NAVY, True, 10, 9, 1.0)
    set_style(doc.styles["Heading 2"], 13, BLUE, True, 12, 6, 1.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Gestinem · Guía para clientes · "), size=8.5, color=MUTED)
    add_page_number(footer)

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_before = Pt(38)
    logo_paragraph.paragraph_format.space_after = Pt(28)
    logo_image = logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(3.35))
    logo_image._inline.docPr.set("descr", "Logotipo corporativo de Gestinem")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("NUEVA APLICACIÓN PARA CLIENTES"), size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(9)
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Gestinem"), size=31, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_font(subtitle.add_run("Guía rápida de acceso y mensajería"), size=16, color=MUTED)
    add_callout(
        doc,
        "Disponible desde el navegador",
        "Las versiones para las tiendas de aplicaciones se publicarán próximamente. Mientras tanto, puede utilizar Gestinem desde Chrome, Edge, Safari o Brave en su móvil, tableta u ordenador. No necesita instalar ningún programa.",
        fill=LIGHT,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(20)
    add_hyperlink(paragraph, URL, URL)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(closing.add_run("Comunicación segura, ordenada y vinculada a su empresa"), size=10, color=MUTED, italic=True)

    new_page(doc, "1. Active su cuenta", "El enlace recibido por correo es personal y caduca en 72 horas.")
    numbers = configure_numbering(doc)
    for text in (
        "Abra el correo de invitación enviado por Gestinem.",
        "Pulse Activar mi cuenta y abrir Gestinem. Se abrirá la aplicación en su navegador.",
        "Cree una contraseña personal de al menos 10 caracteres.",
        "Entre con su correo y la nueva contraseña.",
    ):
        add_list_item(doc, text, numbers)
    add_figure(doc, browser_guide, 6.35, "Acceso a Gestinem desde un navegador actualizado.")
    add_callout(
        doc,
        "Guarde esta dirección",
        f"Para volver a entrar, escriba {URL} en el navegador o guárdela en Favoritos. Cuando las versiones de tienda estén disponibles, Gestinem informará de su publicación.",
        fill=PALE,
        accent=BLUE,
    )

    new_page(doc, "2. Elija el canal adecuado")
    doc.add_paragraph(
        "Cada canal dirige su consulta al equipo correspondiente y mantiene juntos los mensajes y documentos relacionados."
    )
    add_figure(doc, channels_guide, 4.1, "Canales disponibles en la aplicación Gestinem.")
    for heading, detail in (
        ("Laboral", "Nóminas, contratos, altas, bajas y Seguridad Social."),
        ("Contable / Fiscal", "Facturas, impuestos, contabilidad y documentación económica."),
        ("Directo", "Conversación reservada con la persona responsable del despacho."),
    ):
        paragraph = doc.add_paragraph()
        set_font(paragraph.add_run(heading + ". "), color=NAVY, bold=True)
        set_font(paragraph.add_run(detail))

    new_page(doc, "3. Envie mensajes y documentos")
    numbers = configure_numbering(doc)
    for text in (
        "Abra Laboral, Contable / Fiscal o Directo.",
        "Escriba un mensaje breve que explique qué necesita.",
        "Para enviar archivos, pulse el icono de adjuntar y seleccione uno o varios documentos.",
        "Compruebe los nombres de los archivos y pulse Enviar una sola vez.",
    ):
        add_list_item(doc, text, numbers)
    add_callout(
        doc,
        "Formatos habituales",
        "Puede enviar PDF, imágenes, documentos de Word, hojas de cálculo, XML, CSV y ZIP. Utilice nombres claros, por ejemplo Factura_Luz_Agosto_2026.pdf.",
    )
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run("Qué ocurre con los documentos")
    doc.add_paragraph(
        "Los archivos enviados quedan asociados a su conversación. Gestinem puede incorporarlos al expediente o a la gestión documental del despacho sin depender de teléfonos personales."
    )
    add_callout(
        doc,
        "Evite duplicados",
        "Si un archivo tarda unos segundos en aparecer, espere antes de volver a enviarlo.",
        fill="FFF7E8",
        accent="7A5A00",
    )
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run("Recibir avisos")
    doc.add_paragraph(
        "Cuando Gestinem solicite permiso para mostrar notificaciones, pulse Permitir. Puede cambiar esta opción posteriormente desde los ajustes del navegador o del dispositivo."
    )

    new_page(doc, "4. Seguridad y ayuda")
    bullets = configure_numbering(doc, bullet=True)
    for text in (
        "No comparta su contraseña ni el enlace de invitación.",
        "Compruebe que la dirección del navegador empieza por https://app.gestinem.es.",
        "Cierre la sesión si utiliza un ordenador compartido.",
        "Por seguridad, los avisos por correo no incluyen el contenido de sus mensajes.",
    ):
        add_list_item(doc, text, bullets)
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run("Si ha olvidado la contraseña")
    doc.add_paragraph(
        "Pulse ¿Has olvidado tu contraseña? en la pantalla de acceso. Recibirá un enlace seguro para crear una nueva; el enlace caduca en una hora."
    )
    add_callout(
        doc,
        "¿Necesita ayuda?",
        "Si no recibe la invitación, revise la carpeta de correo no deseado. Para problemas de acceso o consultas urgentes, contacte con Gestinem por los medios habituales.",
        fill="EAF6EF",
        accent=GREEN,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(34)
    set_font(
        paragraph.add_run("Gracias por ayudarnos a ofrecerle una atención más rápida, segura y organizada."),
        size=13,
        color=NAVY,
        bold=True,
    )

    doc.core_properties.title = "Gestinem - Guía rápida para clientes"
    doc.core_properties.subject = "Acceso web, canales, documentos, avisos y seguridad"
    doc.core_properties.author = "Gestinem"
    doc.core_properties.keywords = "Gestinem, aplicación Flutter, clientes, mensajería, documentos"
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_COPY.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, OUTPUT_COPY)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
