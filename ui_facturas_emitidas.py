

import os
import calendar
import struct
import tempfile
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import date, datetime

from procesos.facturas_emitidas import generar_emitidas
from utilidades import validar_subcuenta_longitud

IVA_OPCIONES = [21, 10, 4, 0]
IRPF_OPCIONES = [0, 1, 7, 15, 19]


def _to_float(x) -> float:
    try:
        if x is None or x == "":
            return 0.0
        if isinstance(x, (int, float)) and not isinstance(x, bool):
            return float(x)
        s = str(x).strip().replace("\xa0", " ")
        if "." in s and "," in s:
            s = s.replace(".", "").replace(",", ".")
        elif "," in s:
            s = s.replace(",", ".")
        return float(s)
    except Exception:
        return 0.0


def _parse_date_ui(val: str) -> date:
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y"):
        try:
            return datetime.strptime(val.strip(), fmt).date()
        except Exception:
            continue
    return date.today()


def _to_fecha_ui(val: str) -> str:
    if not val:
        return date.today().strftime("%d/%m/%Y")
    try:
        d = _parse_date_ui(str(val))
        return d.strftime("%d/%m/%Y")
    except Exception:
        return date.today().strftime("%d/%m/%Y")


def _to_fecha_ui_or_blank(val: str) -> str:
    if not val:
        return ""
    return _to_fecha_ui(val)


def _round2(x) -> float:
    try:
        return round(float(x), 2)
    except Exception:
        return 0.0


class DatePicker(tk.Toplevel):
    def __init__(self, parent, initial: date | None = None):
        super().__init__(parent)
        self.title("Selecciona fecha")
        self.resizable(False, False)
        self.result = None
        self._current = initial or date.today()
        self._build()
        self.grab_set()
        self.transient(parent)
        self.wait_window(self)

    def _build(self):
        frm = ttk.Frame(self, padding=8)
        frm.pack(fill="both", expand=True)
        nav = ttk.Frame(frm)
        nav.pack(fill="x")
        ttk.Button(nav, text="<", width=3, command=self._prev_month).pack(side=tk.LEFT)
        self.lbl_month = ttk.Label(nav, width=18, anchor="center")
        self.lbl_month.pack(side=tk.LEFT, expand=True)
        ttk.Button(nav, text=">", width=3, command=self._next_month).pack(side=tk.LEFT)

        self.days_frame = ttk.Frame(frm)
        self.days_frame.pack(pady=4)
        for i, dname in enumerate(["Lu", "Ma", "Mi", "Ju", "Vi", "Sa", "Do"]):
            ttk.Label(self.days_frame, text=dname, width=3, anchor="center").grid(row=0, column=i, padx=1, pady=1)

        self._paint_calendar()

    def _paint_calendar(self):
        for w in self.days_frame.grid_slaves():
            info = w.grid_info()
            if int(info.get("row", 0)) > 0:
                w.destroy()
        y, m = self._current.year, self._current.month
        meses = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
        self.lbl_month.config(text=f"{meses[m]} {y}")
        month_days = calendar.monthcalendar(y, m)
        for r, week in enumerate(month_days, start=1):
            for c, day in enumerate(week):
                if day == 0:
                    continue
                btn = ttk.Button(
                    self.days_frame,
                    text=str(day),
                    width=3,
                    command=lambda dd=day: self._select(dd),
                )
                btn.grid(row=r, column=c, padx=1, pady=1)

    def _select(self, day: int):
        try:
            self.result = self._current.replace(day=day)
        except ValueError:
            self.result = None
        self.destroy()

    def _prev_month(self):
        y, m = self._current.year, self._current.month
        if m == 1:
            y -= 1
            m = 12
        else:
            m -= 1
        self._current = self._current.replace(year=y, month=m, day=1)
        self._paint_calendar()

    def _next_month(self):
        y, m = self._current.year, self._current.month
        if m == 12:
            y += 1
            m = 1
        else:
            m += 1
        self._current = self._current.replace(year=y, month=m, day=1)
        self._paint_calendar()


class TerceroFicha(tk.Toplevel):
    def __init__(self, parent, tercero=None):
        super().__init__(parent)
        self.title("Tercero")
        self.resizable(False, False)
        self.result = None
        t = tercero or {}
        fields = [
            ("NIF", "nif", 18),
            ("Nombre", "nombre", 40),
            ("Direccion", "direccion", 40),
            ("CP", "cp", 10),
            ("Poblacion", "poblacion", 28),
            ("Provincia", "provincia", 28),
            ("Telefono", "telefono", 20),
            ("Email", "email", 28),
            ("Contacto", "contacto", 28),
        ]
        self.vars = {}
        for i, (lbl, key, width) in enumerate(fields):
            ttk.Label(self, text=lbl).grid(row=i, column=0, sticky="w", padx=6, pady=3)
            v = tk.StringVar(value=str(t.get(key, "")))
            self.vars[key] = v
            ttk.Entry(self, textvariable=v, width=width).grid(row=i, column=1, sticky="w", padx=6, pady=3)
        ttk.Label(self, text="Tipo").grid(row=len(fields), column=0, sticky="w", padx=6, pady=3)
        self.var_tipo = tk.StringVar(value=t.get("tipo", "cliente"))
        ttk.Combobox(self, textvariable=self.var_tipo, values=["cliente","proveedor","ambos"], state="readonly", width=18).grid(row=len(fields), column=1, sticky="w", padx=6, pady=3)
        btns = ttk.Frame(self)
        btns.grid(row=len(fields)+1, column=0, columnspan=2, pady=6)
        ttk.Button(btns, text="Guardar", style="Primary.TButton", command=self._ok).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Cancelar", command=self.destroy).pack(side=tk.LEFT, padx=4)
        self.grab_set()
        self.transient(parent)
        self.wait_window(self)

    def _ok(self):
        data = {k: v.get().strip() for k, v in self.vars.items()}
        data["tipo"] = self.var_tipo.get() or "cliente"
        self.result = data
        self.destroy()


class TercerosDialog(tk.Toplevel):
    def __init__(self, parent, gestor, codigo_empresa, ejercicio, ndig_plan):
        super().__init__(parent)
        self.title("Terceros")
        self.resizable(True, True)
        self.gestor = gestor
        self.codigo = codigo_empresa
        self.ejercicio = ejercicio
        self.ndig = ndig_plan
        self._build()
        self.grab_set()
        self.transient(parent)
        self.wait_window(self)

    def _build(self):
        frm = ttk.Frame(self, padding=10)
        frm.pack(fill="both", expand=True)

        bar = ttk.Frame(frm)
        bar.pack(fill="x", pady=(0, 6))
        ttk.Button(bar, text="Nuevo", style="Primary.TButton", command=self._nuevo).pack(side=tk.LEFT, padx=4)
        ttk.Button(bar, text="Editar", command=self._editar).pack(side=tk.LEFT, padx=4)
        ttk.Button(bar, text="Eliminar", command=self._eliminar).pack(side=tk.LEFT, padx=4)

        cols = ("nif", "nombre", "tipo", "poblacion")
        self.tv = ttk.Treeview(frm, columns=cols, show="headings", height=12, selectmode="browse")
        self.tv.heading("nif", text="NIF")
        self.tv.column("nif", width=120)
        self.tv.heading("nombre", text="Nombre")
        self.tv.column("nombre", width=240)
        self.tv.heading("tipo", text="Tipo")
        self.tv.column("tipo", width=100)
        self.tv.heading("poblacion", text="Poblacion")
        self.tv.column("poblacion", width=160)
        self.tv.pack(fill="both", expand=True, pady=6)
        self.tv.bind("<<TreeviewSelect>>", lambda e: self._load_subcuentas())

        sub = ttk.LabelFrame(frm, text=f"Subcuentas en empresa {self.codigo}")
        sub.pack(fill="x", pady=(4, 0))
        ttk.Label(sub, text="Subcta cliente").grid(row=0, column=0, sticky="w", padx=6, pady=4)
        ttk.Label(sub, text="Subcta proveedor").grid(row=1, column=0, sticky="w", padx=6, pady=4)
        self.var_sub_cli = tk.StringVar()
        ttk.Entry(sub, textvariable=self.var_sub_cli, width=18).grid(row=0, column=1, sticky="w", padx=6, pady=4)
        self.var_sub_pro = tk.StringVar()
        ttk.Entry(sub, textvariable=self.var_sub_pro, width=18).grid(row=1, column=1, sticky="w", padx=6, pady=4)
        ttk.Button(sub, text="Guardar", style="Primary.TButton", command=self._guardar_sub).grid(row=0, column=2, rowspan=2, padx=6, pady=4)

        self._refresh()

    def _refresh(self):
        self.tv.delete(*self.tv.get_children())
        for t in self.gestor.listar_terceros():
            self.tv.insert(
                "",
                tk.END,
                iid=str(t.get("id")),
                values=(t.get("nif", ""), t.get("nombre", ""), t.get("tipo", "cliente"), t.get("poblacion", "")),
            )
        self.var_sub_cli.set("")
        self.var_sub_pro.set("")

    def _sel_id(self):
        sel = self.tv.selection()
        return sel[0] if sel else None

    def _nuevo(self):
        dlg = TerceroFicha(self)
        if dlg.result:
            tid = self.gestor.upsert_tercero(dlg.result)
            self._refresh()
            self.tv.selection_set(str(tid))

    def _editar(self):
        tid = self._sel_id()
        if not tid:
            return
        ter = next((t for t in self.gestor.listar_terceros() if str(t.get("id")) == str(tid)), None)
        dlg = TerceroFicha(self, ter)
        if dlg.result:
            dlg.result["id"] = tid
            self.gestor.upsert_tercero(dlg.result)
            self._refresh()

    def _eliminar(self):
        tid = self._sel_id()
        if not tid:
            return
        if not messagebox.askyesno("Gest2A3Eco", "Eliminar el tercero seleccionado?"):
            return
        self.gestor.eliminar_tercero(tid)
        self._refresh()

    def _load_subcuentas(self):
        tid = self._sel_id()
        if not tid:
            self.var_sub_cli.set("")
            self.var_sub_pro.set("")
            return
        rel = self.gestor.get_tercero_empresa(self.codigo, tid, self.ejercicio) or {}
        self.var_sub_cli.set(rel.get("subcuenta_cliente", ""))
        self.var_sub_pro.set(rel.get("subcuenta_proveedor", ""))

    def _guardar_sub(self):
        tid = self._sel_id()
        if not tid:
            messagebox.showinfo("Gest2A3Eco", "Selecciona un tercero.")
            return
        sc = self.var_sub_cli.get().strip()
        sp = self.var_sub_pro.get().strip()
        if sc:
            validar_subcuenta_longitud(sc, self.ndig, "subcuenta cliente")
        if sp:
            validar_subcuenta_longitud(sp, self.ndig, "subcuenta proveedor")
        rel = {
            "tercero_id": tid,
            "codigo_empresa": self.codigo,
            "ejercicio": self.ejercicio,
            "subcuenta_cliente": sc,
            "subcuenta_proveedor": sp,
        }
        self.gestor.upsert_tercero_empresa(rel)
        messagebox.showinfo("Gest2A3Eco", "Subcuentas guardadas.")

class FacturaDialog(tk.Toplevel):
    def __init__(self, parent, gestor, codigo_empresa, ejercicio, ndig_plan, factura=None, numero_sugerido=""):
        super().__init__(parent)
        self.title("Factura emitida")
        self.resizable(True, True)
        self.result = None
        self.gestor = gestor
        self.codigo = codigo_empresa
        self.ejercicio = ejercicio
        self.ndig = ndig_plan
        self.factura = dict(factura or {})
        self.factura.setdefault("codigo_empresa", codigo_empresa)
        self.factura.setdefault("ejercicio", ejercicio)
        if "ejercicio" not in self.factura:
            self.factura["ejercicio"] = ejercicio
        if numero_sugerido and not self.factura.get("numero"):
            self.factura["numero"] = numero_sugerido
        f = self.factura

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill="both", expand=True)

        def add_row(label, var, row_idx, width=16, col=0):
            ttk.Label(frm, text=label).grid(row=row_idx, column=col, sticky="w", padx=4, pady=3)
            ttk.Entry(frm, textvariable=var, width=width).grid(row=row_idx, column=col + 1, padx=4, pady=3, sticky="w")

        today = date.today().strftime("%d/%m/%Y")
        self.var_serie = tk.StringVar(value=f.get("serie", ""))
        self.var_numero = tk.StringVar(value=f.get("numero", ""))
        self.var_fecha_asiento = tk.StringVar(value=_to_fecha_ui(f.get("fecha_asiento", today)))
        self.var_fecha_exp = tk.StringVar(value=_to_fecha_ui(f.get("fecha_expedicion", f.get("fecha_asiento", today))))
        self.var_fecha_op = tk.StringVar(value=_to_fecha_ui(f.get("fecha_operacion", today)) if f.get("fecha_operacion") else "")
        self.var_nif = tk.StringVar(value=f.get("nif", ""))
        self.var_nombre = tk.StringVar(value=f.get("nombre", ""))
        self.var_desc = tk.StringVar(value=f.get("descripcion", ""))
        self.var_subcuenta = tk.StringVar(value=f.get("subcuenta_cliente", ""))

        row = 0
        ttk.Label(frm, text="Seleccione Cliente").grid(row=row, column=0, sticky="w", padx=4, pady=3)
        self.var_tercero = tk.StringVar()
        self.cb_tercero = ttk.Combobox(frm, textvariable=self.var_tercero, width=40, state="readonly")
        self.cb_tercero.grid(row=row, column=1, padx=4, pady=3, sticky="w")
        ttk.Button(frm, text="Terceros...", command=self._gestionar_terceros).grid(row=row, column=2, padx=4, pady=3)
        row += 1

        add_row("Serie", self.var_serie, row, width=8, col=0)
        row += 1
        
        add_row("Numero", self.var_numero, row, width=14, col=0)
        row += 1
        
        add_row("Subcuenta cliente", self.var_subcuenta, row, col=0, width=18)
        row += 1

        def add_date_cell(label, var, row_idx):
            ttk.Label(frm, text=label).grid(row=row_idx, column=0, sticky="w", padx=4, pady=3)
            cont = ttk.Frame(frm)
            cont.grid(row=row_idx, column=1, columnspan=2, sticky="w", padx=4, pady=3)
            ttk.Entry(cont, textvariable=var, width=14).pack(side=tk.LEFT)
            ttk.Button(cont, text="Cal", width=4, command=lambda v=var: self._pick_date(v)).pack(side=tk.LEFT, padx=(4, 0))

        add_date_cell("Fecha Factura", self.var_fecha_exp, row)
        #add_date_cell("Fecha Expedicion", self.var_fecha_exp, row, 2)
        #add_date_cell("Fecha Operacion", self.var_fecha_op, row, 3)
        row += 1
        add_row("NIF Cliente", self.var_nif, row)
        row += 1
        add_row("Nombre Cliente", self.var_nombre, row, width=34)
        row += 1
        #add_row("Descripcion", self.var_desc, row, width=40)
        #row += 1

        ttk.Label(frm, text="Lineas").grid(row=row, column=0, columnspan=3, sticky="w", pady=(10, 4))
        row += 1

        self.line_vars = {
            "concepto": tk.StringVar(),
            "unidades": tk.StringVar(),
            "precio": tk.StringVar(),
            "iva": tk.StringVar(),
            "irpf": tk.StringVar(),
        }
        editor = ttk.Frame(frm)
        editor.grid(row=row, column=0, columnspan=3, sticky="ew", padx=4)
        editor.columnconfigure(8, weight=1)
        ttk.Label(editor, text="Concepto").grid(row=0, column=0, padx=4, pady=2, sticky="w")
        ttk.Entry(editor, textvariable=self.line_vars["concepto"], width=26).grid(row=0, column=1, padx=4, pady=2)
        ttk.Label(editor, text="Unidades").grid(row=0, column=2, padx=4, pady=2, sticky="w")
        ttk.Entry(editor, textvariable=self.line_vars["unidades"], width=10).grid(row=0, column=3, padx=4, pady=2)
        ttk.Label(editor, text="Precio").grid(row=0, column=4, padx=4, pady=2, sticky="w")
        ttk.Entry(editor, textvariable=self.line_vars["precio"], width=10).grid(row=0, column=5, padx=4, pady=2)
        ttk.Label(editor, text="IVA %").grid(row=0, column=6, padx=4, pady=2, sticky="w")
        ttk.Combobox(editor, textvariable=self.line_vars["iva"], values=[str(x) for x in IVA_OPCIONES], width=6, state="readonly").grid(row=0, column=7, padx=4, pady=2)
        ttk.Label(editor, text="IRPF %").grid(row=0, column=8, padx=4, pady=2, sticky="w")
        ttk.Combobox(editor, textvariable=self.line_vars["irpf"], values=[str(x) for x in IRPF_OPCIONES], width=6, state="readonly").grid(row=0, column=9, padx=4, pady=2)
        ttk.Button(editor, text="Añadir/Actualizar", style="Primary.TButton", command=self._add_update_linea).grid(row=0, column=10, padx=6)
        ttk.Button(editor, text="Limpiar", command=self._clear_line_editor).grid(row=0, column=11, padx=4)
        row += 1

        self.tv = ttk.Treeview(
            frm,
            columns=("concepto", "unidades", "precio", "base", "pct_iva", "cuota_iva", "pct_irpf", "cuota_irpf"),
            show="headings",
            height=8,
            selectmode="browse",
        )
        headers = {
            "concepto": "Concepto",
            "unidades": "Unid",
            "precio": "P. unit",
            "base": "Base",
            "pct_iva": "% IVA",
            "cuota_iva": "Cuota IVA",
            "pct_irpf": "% IRPF",
            "cuota_irpf": "Ret",
        }
        for c, h in headers.items():
            self.tv.heading(c, text=h)
            self.tv.column(c, width=90 if c == "concepto" else 70, anchor="e" if c != "concepto" else "w")
        self.tv.grid(row=row, column=0, columnspan=3, sticky="nsew", padx=4, pady=4)
        self.tv.bind("<<TreeviewSelect>>", self._on_select_linea)
        row += 1

        bar = ttk.Frame(frm)
        bar.grid(row=row, column=0, columnspan=3, sticky="w", pady=(0, 6))
        ttk.Button(bar, text="Eliminar linea", command=self._del_linea).pack(side=tk.LEFT, padx=4)
        row += 1

        tot = ttk.Frame(frm)
        tot.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(2, 8))
        self.lbl_tot_base = ttk.Label(tot, text="Base: 0.00")
        self.lbl_tot_iva = ttk.Label(tot, text="IVA: 0.00")
        self.lbl_tot_ret = ttk.Label(tot, text="IRPF: 0.00")
        self.lbl_tot_total = ttk.Label(tot, text="Total: 0.00", font=("Segoe UI", 10, "bold"))
        for i, lbl in enumerate([self.lbl_tot_base, self.lbl_tot_iva, self.lbl_tot_ret, self.lbl_tot_total]):
            lbl.grid(row=0, column=i, padx=6)
        row += 1

        btns = ttk.Frame(frm)
        btns.grid(row=row, column=0, columnspan=3, pady=(6, 2))
        ttk.Button(btns, text="Guardar", style="Primary.TButton", command=self._ok).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Cancelar", command=self.destroy).pack(side=tk.LEFT, padx=4)

        self._load_terceros()
        self._preselect_tercero(f.get("tercero_id"))

        for ln in f.get("lineas", []):
            self._insert_linea(ln)
        self._refresh_totales()

        self.grab_set()
        self.transient(parent)
        self.wait_window(self)

    # --- terceros helpers
    def _load_terceros(self):
        self._terceros_cache = self.gestor.listar_terceros()
        disp = [f"{t.get('nombre','')} ({t.get('nif','')})" for t in self._terceros_cache]
        self.cb_tercero["values"] = disp
        self.cb_tercero.bind("<<ComboboxSelected>>", lambda e: self._on_tercero_selected())

    def _preselect_tercero(self, tercero_id):
        if not tercero_id:
            return
        for idx, t in enumerate(self._terceros_cache):
            if str(t.get("id")) == str(tercero_id):
                self.cb_tercero.current(idx)
                self._on_tercero_selected()
                return

    def _gestionar_terceros(self):
        TercerosDialog(self, self.gestor, self.codigo, self.ejercicio, self.ndig)
        self._load_terceros()

    def _on_tercero_selected(self):
        idx = self.cb_tercero.current()
        if idx < 0:
            return
        t = self._terceros_cache[idx]
        self.var_nif.set(t.get("nif", ""))
        self.var_nombre.set(t.get("nombre", ""))
        rel = self.gestor.get_tercero_empresa(self.codigo, t.get("id"), self.ejercicio) or {}
        sc = rel.get("subcuenta_cliente", "")
        if sc:
            self.var_subcuenta.set(sc)

    def _pick_date(self, target_var: tk.StringVar):
        txt = (target_var.get() or "").strip()
        initial = _parse_date_ui(txt) if txt else date.today()
        dlg = DatePicker(self, initial)
        if dlg.result:
            target_var.set(dlg.result.strftime("%d/%m/%Y"))

    # --- line editor
    def _clear_line_editor(self):
        for v in self.line_vars.values():
            v.set("")
        self.tv.selection_remove(self.tv.selection())

    def _line_from_editor(self):
        concepto = self.line_vars["concepto"].get().strip()
        unidades = _to_float(self.line_vars["unidades"].get())
        precio = _to_float(self.line_vars["precio"].get())
        iva = _to_float(self.line_vars["iva"].get())
        irpf = _to_float(self.line_vars["irpf"].get())
        if (
            not concepto
            or unidades == 0
            or precio == 0
            or self.line_vars["iva"].get() == ""
            or self.line_vars["irpf"].get() == ""
        ):
            messagebox.showwarning("Gest2A3Eco", "Completa concepto, unidades, precio, IVA e IRPF.")
            return None
        base = _round2(unidades * precio)
        cuota_iva = _round2(base * iva / 100.0)
        cuota_ret = -abs(_round2(base * irpf / 100.0))
        return {
            "concepto": concepto,
            "unidades": _round2(unidades),
            "precio": _round2(precio),
            "base": _round2(base),
            "pct_iva": _round2(iva),
            "cuota_iva": _round2(cuota_iva),
            "pct_irpf": _round2(irpf),
            "cuota_irpf": _round2(cuota_ret),
            "pct_re": 0.0,
            "cuota_re": 0.0,
        }

    def _add_update_linea(self):
        ln = self._line_from_editor()
        if not ln:
            return
        vals = (
            ln["concepto"],
            f"{ln['unidades']:.2f}",
            f"{ln['precio']:.2f}",
            f"{ln['base']:.2f}",
            f"{ln['pct_iva']:.2f}",
            f"{ln['cuota_iva']:.2f}",
            f"{ln['pct_irpf']:.2f}",
            f"{ln['cuota_irpf']:.2f}",
        )
        sel = self.tv.selection()
        if sel:
            self.tv.item(sel[0], values=vals)
        else:
            self.tv.insert("", tk.END, values=vals)
        self._refresh_totales()
        self._clear_line_editor()

    def _insert_linea(self, ln: dict):
        """Inserta una linea existente (editar/copiar) en la tabla."""
        vals = (
            ln.get("concepto", ""),
            f"{_round2(ln.get('unidades')):.2f}",
            f"{_round2(ln.get('precio')):.2f}",
            f"{_round2(ln.get('base')):.2f}",
            f"{_round2(ln.get('pct_iva')):.2f}",
            f"{_round2(ln.get('cuota_iva')):.2f}",
            f"{_round2(ln.get('pct_irpf')):.2f}",
            f"{_round2(ln.get('cuota_irpf')):.2f}",
        )
        self.tv.insert("", tk.END, values=vals)

    def _on_select_linea(self, event=None):
        sel = self.tv.selection()
        if not sel:
            return
        vals = self.tv.item(sel[0], "values")
        keys = ["concepto", "unidades", "precio", "base", "pct_iva", "cuota_iva", "pct_irpf", "cuota_irpf"]
        for k, v in zip(keys, vals):
            if k in self.line_vars:
                self.line_vars[k].set(v)

    def _del_linea(self):
        sel = self.tv.selection()
        if sel:
            self.tv.delete(sel[0])
            self._refresh_totales()
            self._clear_line_editor()

    def _get_lineas(self):
        out = []
        for iid in self.tv.get_children():
            vals = self.tv.item(iid, "values")
            if not vals:
                continue
            out.append(
                {
                    "concepto": vals[0],
                    "unidades": _round2(_to_float(vals[1])),
                    "precio": _round2(_to_float(vals[2])),
                    "base": _round2(_to_float(vals[3])),
                    "pct_iva": _round2(_to_float(vals[4])),
                    "cuota_iva": _round2(_to_float(vals[5])),
                    "pct_irpf": _round2(_to_float(vals[6])),
                    "cuota_irpf": _round2(_to_float(vals[7])),
                    "pct_re": 0.0,
                    "cuota_re": 0.0,
                }
            )
        return out

    def _refresh_totales(self):
        base = iva = ret = 0.0
        for ln in self._get_lineas():
            base += ln["base"]
            iva += ln["cuota_iva"]
            ret += ln["cuota_irpf"]
        base = _round2(base)
        iva = _round2(iva)
        ret = _round2(ret)
        total = _round2(base + iva + ret)
        self.lbl_tot_base.config(text=f"Base: {base:.2f}")
        self.lbl_tot_iva.config(text=f"IVA: {iva:.2f}")
        self.lbl_tot_ret.config(text=f"IRPF: {ret:.2f}")
        self.lbl_tot_total.config(text=f"Total: {total:.2f}")

    def _ok(self):
        lineas = self._get_lineas()
        if not lineas:
            messagebox.showerror("Gest2A3Eco", "Añade al menos una linea.")
            return
        if not self.var_numero.get().strip():
            messagebox.showerror("Gest2A3Eco", "Numero de factura vacio.")
            return
        sc = self.var_subcuenta.get().strip()
        if sc:
            validar_subcuenta_longitud(sc, self.ndig, "subcuenta cliente")
        tercero_id = None
        idx = self.cb_tercero.current()
        if idx >= 0:
            tercero_id = self._terceros_cache[idx].get("id")
        fecha_common = self.var_fecha_exp.get().strip()
        self.result = {
            "id": self.factura.get("id"),
            "codigo_empresa": self.codigo,
            "ejercicio": self.ejercicio,
            "tercero_id": tercero_id,
            "serie": self.var_serie.get().strip(),
            "numero": self.var_numero.get().strip(),
            "numero_largo_sii": self.factura.get("numero_largo_sii", ""),
            "fecha_asiento": fecha_common,
            "fecha_expedicion": fecha_common,
            "fecha_operacion": fecha_common,
            "nif": self.var_nif.get().strip(),
            "nombre": self.var_nombre.get().strip(),
            "descripcion": self.var_desc.get().strip(),
            "subcuenta_cliente": sc,
            "lineas": lineas,
            "generada": self.factura.get("generada", False),
            "fecha_generacion": self.factura.get("fecha_generacion", ""),
        }
        self.destroy()


class UIFacturasEmitidas(ttk.Frame):
    def __init__(self, master, gestor, codigo_empresa, ejercicio, nombre_empresa):
        super().__init__(master)
        self.gestor = gestor
        self.codigo = codigo_empresa
        self.ejercicio = ejercicio
        self.nombre = nombre_empresa
        base = {
            "nombre": nombre_empresa,
            "digitos_plan": 8,
            "serie_emitidas": "A",
            "siguiente_num_emitidas": 1,
            "ejercicio": "",
            "cif": "",
            "direccion": "",
            "poblacion": "",
            "provincia": "",
            "cp": "",
            "telefono": "",
            "email": "",
            "logo_path": "",
        }
        emp_conf = gestor.get_empresa(codigo_empresa, ejercicio) or {}
        base.update(emp_conf)
        self.empresa_conf = base
        self._build()

    # ------------------- UI -------------------
    def _build(self):
        ttk.Label(self, text=f"Facturas emitidas de {self.nombre} ({self.codigo} · {self.ejercicio})", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10, pady=8)

        top = ttk.Frame(self)
        top.pack(fill="x", padx=10)
        ttk.Button(top, text="Nueva", style="Primary.TButton", command=self._nueva).pack(side=tk.LEFT, padx=4)
        ttk.Button(top, text="Editar", command=self._editar).pack(side=tk.LEFT, padx=4)
        ttk.Button(top, text="Copiar", command=self._copiar).pack(side=tk.LEFT, padx=4)
        ttk.Button(top, text="Eliminar", command=self._eliminar).pack(side=tk.LEFT, padx=4)
        ttk.Button(top, text="Terceros", command=self._terceros).pack(side=tk.LEFT, padx=12)
        ttk.Button(top, text="Exportar PDF", command=self._export_pdf).pack(side=tk.LEFT, padx=4)

        self.tv = ttk.Treeview(
            self,
            columns=("serie", "numero", "fecha", "cliente", "total", "generada", "fecha_gen"),
            show="headings",
            selectmode="extended",
            height=12,
        )
        cols = [
            ("serie", "Serie", 80, "w"),
            ("numero", "Numero", 120, "w"),
            ("fecha", "Fecha", 100, "w"),
            ("cliente", "Cliente", 240, "w"),
            ("total", "Total", 100, "e"),
            ("generada", "Generada", 90, "center"),
            ("fecha_gen", "Fecha gen.", 110, "w"),
        ]
        for c, h, w, align in cols:
            self.tv.heading(c, text=h)
            self.tv.column(c, width=w, anchor=align)
        self.tv.pack(fill="both", expand=True, padx=10, pady=8)

        bottom = ttk.Frame(self)
        bottom.pack(fill="x", padx=10, pady=6)
        ttk.Label(bottom, text="Plantilla emitidas:").pack(side=tk.LEFT)
        self.cb_plantilla = ttk.Combobox(bottom, width=40, state="readonly")
        self.cb_plantilla.pack(side=tk.LEFT, padx=6)
        ttk.Button(bottom, text="Generar Suenlace.dat", style="Primary.TButton", command=self._generar).pack(side=tk.RIGHT)

        self._refresh_plantillas()
        self._refresh_facturas()

    # ------------------- Datos -------------------
    def _refresh_plantillas(self):
        pls = [p.get("nombre") for p in self.gestor.listar_emitidas(self.codigo, self.ejercicio)]
        self.cb_plantilla["values"] = pls
        if pls:
            self.cb_plantilla.current(0)

    def _compute_total(self, fac: dict) -> float:
        total = 0.0
        for ln in fac.get("lineas", []):
            total += (
                _to_float(ln.get("base"))
                + _to_float(ln.get("cuota_iva"))
                + _to_float(ln.get("cuota_re"))
                + _to_float(ln.get("cuota_irpf"))
            )
        return _round2(total)

    def _refresh_facturas(self):
        self.tv.delete(*self.tv.get_children())
        for fac in self.gestor.listar_facturas_emitidas(self.codigo, self.ejercicio):
            total = self._compute_total(fac)
            self.tv.insert(
                "",
                tk.END,
                iid=str(fac.get("id")),
                values=(
                    fac.get("serie", ""),
                    fac.get("numero", ""),
                    _to_fecha_ui_or_blank(fac.get("fecha_asiento", "")),
                    fac.get("nombre", ""),
                    f"{total:.2f}",
                    "Si" if fac.get("generada") else "No",
                    fac.get("fecha_generacion", ""),
                ),
            )

    def _selected_ids(self):
        sel = list(self.tv.selection())
        if sel:
            return sel
        focus = self.tv.focus()
        return [focus] if focus else []

    def _serie(self):
        return str(self.empresa_conf.get("serie_emitidas", "A") or "A")

    def _siguiente_num(self):
        try:
            return int(self.empresa_conf.get("siguiente_num_emitidas", 1))
        except Exception:
            return 1

    def _proximo_numero(self):
        return f"{self._serie()}{self._siguiente_num():06d}"

    def _incrementar_numeracion(self):
        self.empresa_conf["siguiente_num_emitidas"] = self._siguiente_num() + 1
        self.gestor.upsert_empresa(self.empresa_conf)

    def _nueva(self):
        sugerido = self._proximo_numero()
        dlg = FacturaDialog(
            self,
            self.gestor,
            self.codigo,
            self.ejercicio,
            int(self.empresa_conf.get("digitos_plan", 8)),
            {"codigo_empresa": self.codigo, "ejercicio": self.ejercicio, "serie": self._serie(), "numero": sugerido},
            numero_sugerido=sugerido,
        )
        if dlg.result:
            dlg.result["generada"] = False
            dlg.result["fecha_generacion"] = ""
            dlg.result["ejercicio"] = self.ejercicio
            dlg.result["codigo_empresa"] = self.codigo
            self.gestor.upsert_factura_emitida(dlg.result)
            self._incrementar_numeracion()
            self._refresh_facturas()

    def _editar(self):
        sel = self._selected_ids()
        if not sel:
            messagebox.showinfo("Gest2A3Eco", "Selecciona una factura.")
            return
        fac = next((f for f in self.gestor.listar_facturas_emitidas(self.codigo, self.ejercicio) if str(f.get("id")) == str(sel[0])), None)
        if not fac:
            return
        dlg = FacturaDialog(self, self.gestor, self.codigo, self.ejercicio, int(self.empresa_conf.get("digitos_plan", 8)), fac)
        if dlg.result:
            self.gestor.upsert_factura_emitida(dlg.result)
            self._refresh_facturas()

    def _copiar(self):
        sel = self._selected_ids()
        if not sel:
            return
        fac = next((f for f in self.gestor.listar_facturas_emitidas(self.codigo, self.ejercicio) if str(f.get("id")) == str(sel[0])), None)
        if not fac:
            return
        nuevo = dict(fac)
        nuevo.pop("id", None)
        nuevo["numero"] = self._proximo_numero()
        nuevo["serie"] = self._serie()
        nuevo["generada"] = False
        nuevo["fecha_generacion"] = ""
        dlg = FacturaDialog(
            self,
            self.gestor,
            self.codigo,
            self.ejercicio,
            int(self.empresa_conf.get("digitos_plan", 8)),
            nuevo,
            numero_sugerido=nuevo["numero"],
        )
        if dlg.result:
            self.gestor.upsert_factura_emitida(dlg.result)
            self._incrementar_numeracion()
            self._refresh_facturas()

    def _eliminar(self):
        sel = self._selected_ids()
        if not sel:
            return
        if not messagebox.askyesno("Gest2A3Eco", "Eliminar las facturas seleccionadas?"):
            return
        for fid in sel:
            self.gestor.eliminar_factura_emitida(self.codigo, fid, self.ejercicio)
        self._refresh_facturas()

    def _terceros(self):
        TercerosDialog(self, self.gestor, self.codigo, self.ejercicio, int(self.empresa_conf.get("digitos_plan", 8)))

    # ------------------- Exportar PDF -------------------
    def _docx_template_path(self) -> str:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantilla.docx")

    def _docx_placeholder_map(self, fac: dict):
        cli = self._cliente_factura(fac)
        tot = self._totales_factura(fac)
        return {
            "{{empresa_nombre}}": self.empresa_conf.get("nombre", "") or self.nombre,
            "{{empresa_cif}}": self.empresa_conf.get("cif", ""),
            "{{empresa_direccion}}": self.empresa_conf.get("direccion", ""),
            "{{empresa_cp}}": self.empresa_conf.get("cp", ""),
            "{{empresa_poblacion}}": self.empresa_conf.get("poblacion", ""),
            "{{empresa_provincia}}": self.empresa_conf.get("provincia", ""),
            "{{empresa_telefono}}": self.empresa_conf.get("telefono", ""),
            "{{empresa_email}}": self.empresa_conf.get("email", ""),
            "{{factura_serie}}": fac.get("serie", ""),
            "{{factura_numero}}": fac.get("numero", ""),
            "{{factura_fecha}}": _to_fecha_ui_or_blank(fac.get("fecha_expedicion") or fac.get("fecha_asiento", "")),
            "{{factura_fecha_operacion}}": _to_fecha_ui_or_blank(fac.get("fecha_operacion", "")),
            "{{factura_descripcion}}": fac.get("descripcion", ""),
            "{{cliente_nombre}}": cli.get("nombre", ""),
            "{{cliente_nif}}": cli.get("nif", ""),
            "{{cliente_direccion}}": cli.get("direccion", ""),
            "{{cliente_cp}}": cli.get("cp", ""),
            "{{cliente_poblacion}}": cli.get("poblacion", ""),
            "{{cliente_provincia}}": cli.get("provincia", ""),
            "{{cliente_telefono}}": cli.get("telefono", ""),
            "{{cliente_email}}": cli.get("email", ""),
            "{{total_base}}": f"{tot['base']:.2f}",
            "{{total_iva}}": f"{tot['iva']:.2f}",
            "{{total_irpf}}": f"{tot['ret']:.2f}",
            "{{total_factura}}": f"{tot['total']:.2f}",
            "{{observaciones}}": fac.get("descripcion", ""),
            "{{pago_metodo}}": "",
            "{{pago_banco}}": "",
            "{{pago_iban}}": "",
            "{{pago_vencimiento}}": "",
        }

    def _replace_text_runs(self, container, mapping: dict):
        def replace_paragraph(p):
            txt = p.text
            changed = False
            for key, val in mapping.items():
                if key in txt:
                    txt = txt.replace(key, val)
                    changed = True
            if not changed:
                return
            # Reescribir en un solo run para evitar splits de Word
            while p.runs:
                p.runs[0].clear()
                p._p.remove(p.runs[0]._r)
            run = p.add_run(txt)
            return

        for p in getattr(container, "paragraphs", []):
            replace_paragraph(p)
        for tbl in getattr(container, "tables", []):
            for row in tbl.rows:
                for cell in row.cells:
                    self._replace_text_runs(cell, mapping)

    def _linea_placeholder_map(self, ln: dict):
        total_linea = (
            _to_float(ln.get("base"))
            + _to_float(ln.get("cuota_iva"))
            + _to_float(ln.get("cuota_re"))
            + _to_float(ln.get("cuota_irpf"))
        )
        return {
            "{{linea_concepto}}": str(ln.get("concepto", "")),
            "{{linea_unidades}}": f"{_to_float(ln.get('unidades')):.2f}",
            "{{linea_precio}}": f"{_to_float(ln.get('precio')):.2f}",
            "{{linea_pct_iva}}": f"{_to_float(ln.get('pct_iva')):.2f}",
            "{{linea_cuota_iva}}": f"{_to_float(ln.get('cuota_iva')):.2f}",
            "{{linea_pct_irpf}}": f"{_to_float(ln.get('pct_irpf')):.2f}",
            "{{linea_cuota_irpf}}": f"{_to_float(ln.get('cuota_irpf')):.2f}",
            "{{linea_base}}": f"{_to_float(ln.get('base')):.2f}",
            "{{linea_total}}": f"{total_linea:.2f}",
        }

    def _fill_lineas_table(self, doc, lineas: list):
        placeholder = "{{linea_concepto}}"
        for table in doc.tables:
            template_row = None
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                if placeholder in row_text:
                    template_row = row
                    break
            if not template_row:
                continue
            template_texts = [cell.text for cell in template_row.cells]
            if not lineas:
                self._replace_text_runs(template_row, self._linea_placeholder_map({}))
                return True
            for idx, ln in enumerate(lineas):
                target_row = template_row if idx == 0 else table.add_row()
                if idx > 0:
                    for c_idx, txt in enumerate(template_texts):
                        target_row.cells[c_idx].text = txt
                self._replace_text_runs(target_row, self._linea_placeholder_map(ln))
            return True
        return False

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        try:
            import win32com.client  # type: ignore
        except Exception:
            return False
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
            doc.Close(False)
            word.Quit()
            return True
        except Exception:
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            return False

    def _export_pdf_with_template(self, fac: dict, save_path: str):
        tpl_path = self._docx_template_path()
        if not os.path.exists(tpl_path):
            return False, "No se encuentra plantilla.docx en la carpeta del programa."
        try:
            from docx import Document  # type: ignore
        except Exception:
            return False, "Falta la libreria python-docx (instala requirements.txt)."
        try:
            doc = Document(tpl_path)
        except Exception:
            return False, "No se pudo abrir plantilla.docx (esta corrupto?)."
        mapping = self._docx_placeholder_map(fac)
        self._replace_text_runs(doc, mapping)
        self._fill_lineas_table(doc, fac.get("lineas", []))
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        try:
            doc.save(tmp.name)
            ok = self._convert_docx_to_pdf(tmp.name, save_path)
            if not ok:
                return False, "Necesitas Word con pywin32 para convertir a PDF."
            return True, ""
        finally:
            try:
                os.remove(tmp.name)
            except Exception:
                pass

    def _pdf_escape(self, text: str) -> str:
        return str(text or "").replace("\\", "\\\\").replace("(", "[").replace(")", "]")

    def _totales_factura(self, fac: dict):
        base = iva = re = ret = 0.0
        for ln in fac.get("lineas", []):
            base += _to_float(ln.get("base"))
            iva += _to_float(ln.get("cuota_iva"))
            re += _to_float(ln.get("cuota_re"))
            ret += _to_float(ln.get("cuota_irpf"))
        total = base + iva + re + ret
        return {
            "base": _round2(base),
            "iva": _round2(iva),
            "re": _round2(re),
            "ret": _round2(ret),
            "total": _round2(total),
        }

    def _cliente_factura(self, fac: dict):
        cli = next(
            (
                t
                for t in self.gestor.listar_terceros()
                if str(t.get("id")) == str(fac.get("tercero_id"))
            ),
            {},
        )
        return {
            "nombre": fac.get("nombre") or cli.get("nombre", ""),
            "nif": fac.get("nif") or cli.get("nif", ""),
            "direccion": cli.get("direccion", ""),
            "cp": cli.get("cp", ""),
            "poblacion": cli.get("poblacion", ""),
            "provincia": cli.get("provincia", ""),
            "telefono": cli.get("telefono", ""),
            "email": cli.get("email", ""),
        }

    def _logo_jpeg(self):
        path = str(self.empresa_conf.get("logo_path") or "").strip()
        if not path or not os.path.exists(path):
            return None
        if not path.lower().endswith((".jpg", ".jpeg")):
            return None
        try:
            with open(path, "rb") as f:
                data = f.read()
            # Buscar marcador SOF para dimensiones
            i = 0
            while i < len(data) - 9:
                if data[i] != 0xFF:
                    i += 1
                    continue
                marker = data[i + 1]
                if marker in (0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF):
                    if i + 9 >= len(data):
                        break
                    h, w = struct.unpack(">HH", data[i + 5 : i + 9])
                    comp = data[i + 9] if i + 9 < len(data) else 3
                    return {"data": data, "w": w, "h": h, "components": comp}
                else:
                    if i + 4 >= len(data):
                        break
                    seg_len = struct.unpack(">H", data[i + 2 : i + 4])[0]
                    i += seg_len + 2
        except Exception:
            return None
        return None

    def _factura_pdf(self, path: str, fac: dict):
        cliente = self._cliente_factura(fac)
        tot = self._totales_factura(fac)
        obs = fac.get("descripcion") or ""

        def t(x, y, txt, size=11, bold=False):
            font = "F2" if bold else "F1"
            return f"BT /{font} {size} Tf 1 0 0 1 {x} {y} Tm ({self._pdf_escape(txt)}) Tj ET\n"

        y = 800
        body = []
        logo = self._logo_jpeg()
        logo_cmd = ""
        if logo:
            disp_w = 120
            disp_h = max(40, int(logo["h"] * disp_w / max(logo["w"], 1)))
            logo_cmd = f"q {disp_w} 0 0 {disp_h} 50 {y - disp_h + 10} cm /Im1 Do Q\n"
            y -= disp_h + 10

        body.append(t(50, y, "Factura emitida", 16, True))
        y -= 24

        # Bloque emisor (izquierda)
        emitter = self.empresa_conf
        body.append(t(50, y, emitter.get("nombre") or self.nombre, 12, True)); y -= 14
        body.append(t(50, y, f"CIF: {emitter.get('cif','')}", 10)); y -= 12
        dir_line = ", ".join(filter(None, [emitter.get("direccion"), emitter.get("cp"), emitter.get("poblacion")]))
        if dir_line:
            body.append(t(50, y, dir_line, 10)); y -= 12
        prov_line = ", ".join(filter(None, [emitter.get("provincia"), emitter.get("telefono"), emitter.get("email")]))
        if prov_line:
            body.append(t(50, y, prov_line, 10)); y -= 12
        info_eje = emitter.get("ejercicio", "")
        body.append(t(50, y, f"Ejercicio: {info_eje}  Serie: {fac.get('serie','') or '-'}", 10)); y -= 18

        # Bloque cliente (derecha en recuadro)
        box_w, box_h = 260, 70
        box_x, box_y = 320, y
        body.append(f"q 1 w {box_x} {box_y - box_h} {box_w} {box_h} re S Q\n")
        body.append(t(box_x + 8, box_y - 14, "Cliente", 12, True))
        body.append(t(box_x + 8, box_y - 28, f"{cliente.get('nombre','')}", 10))
        body.append(t(box_x + 8, box_y - 42, f"NIF: {cliente.get('nif','')}", 10))
        addr = ", ".join(filter(None, [cliente.get("direccion"), cliente.get("cp"), cliente.get("poblacion")]))
        if addr:
            body.append(t(box_x + 8, box_y - 56, addr, 10))
        contacto = ", ".join(filter(None, [cliente.get("provincia"), cliente.get("telefono"), cliente.get("email")]))
        if contacto:
            body.append(t(box_x + 8, box_y - 70, contacto, 10))
        y = min(y, box_y - box_h - 12)

        # Línea de metadatos factura
        y -= 4
        body.append(t(50, y, f"Factura: {fac.get('serie','')}-{fac.get('numero','')}", 11, True))
        body.append(t(260, y, f"Fecha: {_to_fecha_ui_or_blank(fac.get('fecha_expedicion') or fac.get('fecha_asiento',''))}", 11))
        if obs:
            y -= 12
            body.append(t(50, y, f"Observaciones: {obs}", 10))
        y -= 16

        headers = [
            ("Concepto", 50),
            ("Unid", 270),
            ("P. unit", 320),
            ("Base", 380),
            ("% IVA", 440),
            ("Cuota IVA", 490),
            ("% IRPF", 545),
        ]
        for txt, x in headers:
            body.append(t(x, y, txt, 11, True))
        y -= 12

        for ln in fac.get("lineas", []):
            body.append(t(50, y, str(ln.get("concepto", ""))[:42], 10))
            body.append(t(270, y, f"{_to_float(ln.get('unidades')):.2f}", 10))
            body.append(t(320, y, f"{_to_float(ln.get('precio')):.2f}", 10))
            body.append(t(380, y, f"{_to_float(ln.get('base')):.2f}", 10))
            body.append(t(440, y, f"{_to_float(ln.get('pct_iva')):.2f}%", 10))
            body.append(t(490, y, f"{_to_float(ln.get('cuota_iva')):.2f}", 10))
            body.append(t(545, y, f"{_to_float(ln.get('pct_irpf')):.2f}%", 10))
            y -= 12

        y -= 16
        body.append(t(360, y, "Base imponible:", 11, True))
        body.append(t(500, y, f"{tot['base']:.2f}", 11))
        y -= 14
        body.append(t(360, y, "IVA:", 11, True))
        body.append(t(500, y, f"{tot['iva']:.2f}", 11))
        y -= 14
        if abs(tot["re"]) > 0.001:
            body.append(t(360, y, "Recargo Eq.:", 11, True))
            body.append(t(500, y, f"{tot['re']:.2f}", 11))
            y -= 14
            body.append(t(360, y, "IRPF:", 11, True))
            body.append(t(500, y, f"{tot['ret']:.2f}", 11))
            y -= 16
        body.append(t(360, y, "Total factura:", 12, True))
        body.append(t(500, y, f"{tot['total']:.2f}", 12, True))

        content_parts = []
        if logo_cmd:
            content_parts.append(logo_cmd)
        content_parts.extend(body)
        stream = "".join(content_parts).encode("latin-1", "ignore")
        objs = []
        objs.append("1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
        objs.append("2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")

        # Fuentes
        objs.append(None)  # placeholder for page, lo rellenamos mas abajo
        objs.append(f"4 0 obj << /Length {len(stream)} >> stream\n".encode() + stream + b"endstream\nendobj\n")
        objs.append("5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
        objs.append("6 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >> endobj\n")

        xobject_ref = ""
        if logo:
            color_space = "/DeviceGray" if logo["components"] == 1 else "/DeviceRGB"
            objs.append(
                f"7 0 obj << /Type /XObject /Subtype /Image /Width {logo['w']} /Height {logo['h']} "
                f"/ColorSpace {color_space} /BitsPerComponent 8 /Filter /DCTDecode /Length {len(logo['data'])} >> stream\n".encode()
                + logo["data"]
                + b"endstream\nendobj\n"
            )
            xobject_ref = "/XObject << /Im1 7 0 R >>"

        res_parts = ["/Font << /F1 5 0 R /F2 6 0 R >>"]
        if xobject_ref:
            res_parts.append(xobject_ref)
        page_res = "/Resources << " + " ".join(res_parts) + " >>"
        objs[2] = (
            "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            f"/Contents 4 0 R {page_res} >> endobj\n"
        )
        with open(path, "wb") as f:
            f.write(b"%PDF-1.4\n")
            offsets = []
            for ob in objs:
                ob_bytes = ob if isinstance(ob, bytes) else ob.encode("latin-1")
                offsets.append(f.tell())
                f.write(ob_bytes)
            xref_pos = f.tell()
            f.write(f"xref\n0 {len(objs)+1}\n".encode())
            f.write(b"0000000000 65535 f \n")
            for off in offsets:
                f.write(f"{off:010d} 00000 n \n".encode())
            f.write(f"trailer << /Size {len(objs)+1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode())

    def _export_pdf(self):
        sel = self._selected_ids()
        if not sel:
            messagebox.showinfo("Gest2A3Eco", "Selecciona una factura.")
            return
        fac = next((f for f in self.gestor.listar_facturas_emitidas(self.codigo, self.ejercicio) if str(f.get("id")) == str(sel[0])), None)
        if not fac:
            return

        save_path = filedialog.asksaveasfilename(
            title="Exportar PDF",
            defaultextension=".pdf",
            initialfile=f"Factura_{fac.get('numero','')}.pdf",
            filetypes=[("PDF", "*.pdf")],
        )
        if not save_path:
            return
        try:
            ok_tpl, reason = self._export_pdf_with_template(fac, save_path)
            if ok_tpl:
                messagebox.showinfo("Gest2A3Eco", f"PDF generado con plantilla.docx:\n{save_path}")
                return
            if reason:
                messagebox.showwarning("Gest2A3Eco", f"No se uso la plantilla.docx:\n{reason}\nSe intentara el PDF basico.")
        except Exception:
            pass
        try:
            self._factura_pdf(save_path, fac)
            messagebox.showinfo("Gest2A3Eco", f"PDF generado (modo basico):\n{save_path}")
        except Exception as e:
            messagebox.showerror("Gest2A3Eco", f"No se pudo generar el PDF:\n{e}")

    def _factura_to_rows(self, fac: dict):
        rows = []
        base_row = {
            "Serie": fac.get("serie", ""),
            "Numero Factura": fac.get("numero", ""),
            "Numero Factura Largo SII": fac.get("numero_largo_sii", ""),
            "Fecha Asiento": fac.get("fecha_asiento", ""),
            "Fecha Expedicion": fac.get("fecha_expedicion") or fac.get("fecha_asiento", ""),
            "Fecha Operacion": fac.get("fecha_operacion", ""),
            "Descripcion Factura": fac.get("descripcion", ""),
            "NIF Cliente Proveedor": fac.get("nif", ""),
            "Nombre Cliente Proveedor": fac.get("nombre", ""),
        }
        if fac.get("subcuenta_cliente"):
            base_row["Cuenta Cliente Proveedor"] = fac.get("subcuenta_cliente")
        for ln in fac.get("lineas", []):
            r = dict(base_row)
            r.update(
                {
                    "Descripcion Linea": ln.get("concepto", "") or fac.get("descripcion", ""),
                    "Base": _round2(_to_float(ln.get("base"))),
                    "Cuota IVA": _round2(_to_float(ln.get("cuota_iva"))),
                    "Porcentaje IVA": _round2(_to_float(ln.get("pct_iva"))),
                    "Porcentaje Recargo Equivalencia": _round2(_to_float(ln.get("pct_re"))),
                    "Cuota Recargo Equivalencia": _round2(_to_float(ln.get("cuota_re"))),
                    "Porcentaje Retencion IRPF": _round2(_to_float(ln.get("pct_irpf"))),
                    "Cuota Retencion IRPF": _round2(_to_float(ln.get("cuota_irpf"))),
                }
            )
            rows.append(r)
        return rows

    def _generar(self):
        sel = self._selected_ids()
        if not sel:
            messagebox.showwarning("Gest2A3Eco", "Selecciona al menos una factura.")
            return
        pl_name = (self.cb_plantilla.get() or "").strip()
        if not pl_name:
            messagebox.showwarning("Gest2A3Eco", "Selecciona una plantilla de emitidas.")
            return
        plantilla = next((p for p in self.gestor.listar_emitidas(self.codigo, self.ejercicio) if p.get("nombre") == pl_name), None)
        if not plantilla:
            messagebox.showerror("Gest2A3Eco", "Plantilla no encontrada.")
            return

        facturas_sel = []
        rows = []
        for fid in sel:
            fac = next((f for f in self.gestor.listar_facturas_emitidas(self.codigo, self.ejercicio) if str(f.get("id")) == str(fid)), None)
            if fac:
                facturas_sel.append(fac)
                rows.extend(self._factura_to_rows(fac))

        ya_generadas = [f for f in facturas_sel if f.get("generada")]
        if ya_generadas:
            nums = ", ".join(f"{f.get('serie','')}-{f.get('numero','')}" for f in ya_generadas)
            if not messagebox.askyesno(
                "Gest2A3Eco",
                f"Las facturas {nums} ya están marcadas como generadas.\n¿Generar suenlace de todas formas?",
            ):
                return

        if not rows:
            messagebox.showwarning("Gest2A3Eco", "No hay lineas para generar.")
            return

        ndig = int(self.empresa_conf.get("digitos_plan", 8))
        registros = generar_emitidas(rows, plantilla, str(self.codigo), ndig)
        if not registros:
            messagebox.showwarning("Gest2A3Eco", "No se generaron registros.")
            return

        save_path = filedialog.asksaveasfilename(
            title="Guardar fichero suenlace.dat",
            defaultextension=".dat",
            initialfile=f"E{self.codigo}.dat",
            filetypes=[("Ficheros DAT", "*.dat")],
        )
        if not save_path:
            return
        with open(save_path, "w", encoding="latin-1", newline="") as f:
            f.writelines(registros)
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M")
        self.gestor.marcar_facturas_emitidas_generadas(self.codigo, sel, fecha_gen, self.ejercicio)
        self._refresh_facturas()
        messagebox.showinfo("Gest2A3Eco", f"Fichero generado:\n{save_path}")
