"""Plantillas Word configurables para generar documentos de firma."""
from __future__ import annotations

import hashlib
import re
import shutil
import uuid
from datetime import date, datetime
from pathlib import Path

from docx import Document

from procesos.facturas_word import convert_docx_to_pdf, render_docx
from utils.utilidades import (
    get_document_repository_dir,
    get_word_templates_dir,
    get_word_templates_subdir,
)
from utils.validaciones import normalizar_nif_cif, validar_nif_cif_nie


TIPOS_CARPETA = ("facturas", "albaranes", "firmas")
ORIGENES_CAMPO = ("empresa", "tercero", "sistema", "fijo", "manual")
TIPOS_CAMPO = ("texto", "texto_largo", "fecha", "numero", "nif", "email", "telefono", "booleano")
ORIGENES_FIRMANTE = ("empresa", "tercero", "gestor", "manual")
CAMPOS_POR_ORIGEN = {
    "empresa": ("codigo", "nombre", "cif", "direccion", "cp", "poblacion", "provincia", "pais", "telefono", "email", "naf", "responsable"),
    "tercero": ("nombre", "nombre_legal", "nombre_comercial", "nif", "direccion", "cp", "poblacion", "provincia", "pais", "telefono", "email", "contacto"),
    "sistema": ("fecha_actual", "anio_actual", "usuario_nombre"),
}

_ALIAS_CAMPO_ORIGEN = {
    "empresa": {
        "nombre_legal": "nombre", "nombre_comercial": "nombre",
        "nif": "cif", "dni": "cif", "nif_cif": "cif",
        "codigo_postal": "cp", "municipio": "poblacion", "localidad": "poblacion",
        "correo": "email",
    },
    "tercero": {
        "cif": "nif", "dni": "nif", "nif_cif": "nif",
        "codigo_postal": "cp", "municipio": "poblacion", "localidad": "poblacion",
        "correo": "email",
    },
}

_TAG = re.compile(r"{{\s*([A-Za-z_][A-Za-z0-9_]*)\s*}}")
_EMAIL = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_SAFE = re.compile(r"[^A-Za-z0-9_. -]+")


class PlantillasFirmaService:
    def __init__(self, gestor):
        self.gestor = gestor

    @staticmethod
    def asegurar_subcarpetas() -> dict[str, Path]:
        return {tipo: get_word_templates_subdir(tipo) for tipo in (*TIPOS_CARPETA, "tramites_dgt")}

    @staticmethod
    def listar_raiz_para_organizar() -> list[Path]:
        root = Path(get_word_templates_dir())
        return sorted((p for p in root.glob("*.docx") if p.is_file()), key=lambda p: p.name.lower())

    @classmethod
    def copiar_clasificadas(cls, clasificacion: dict[str, str]) -> dict:
        """Copia sin borrar; los conflictos diferentes quedan pendientes."""
        root = Path(get_word_templates_dir())
        resumen = {"copiados": [], "identicos": [], "conflictos": [], "ignorados": []}
        for nombre, tipo in dict(clasificacion or {}).items():
            source = root / Path(nombre).name
            if tipo not in TIPOS_CARPETA or not source.is_file():
                resumen["ignorados"].append(source.name)
                continue
            target = get_word_templates_subdir(tipo) / source.name
            if target.exists():
                bucket = "identicos" if cls.sha256(target) == cls.sha256(source) else "conflictos"
                resumen[bucket].append(source.name)
                continue
            shutil.copy2(source, target)
            resumen["copiados"].append(source.name)
        return resumen

    @staticmethod
    def detectar_campos(path: str | Path) -> list[str]:
        path = Path(path)
        if not path.is_file() or path.suffix.lower() != ".docx":
            raise ValueError("La plantilla debe ser un archivo DOCX valido.")
        doc = Document(str(path))
        textos: list[str] = []

        def extraer_bloque(contenedor) -> None:
            for paragraph in contenedor.paragraphs:
                textos.append(paragraph.text)
            for table in contenedor.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extraer_bloque(cell)

        extraer_bloque(doc)
        for section in doc.sections:
            extraer_bloque(section.header)
            extraer_bloque(section.footer)
        texto = "\n".join(textos)
        if "{%" in texto or "{#" in texto:
            raise ValueError("Las plantillas de firma no admiten bucles, condicionales ni comentarios Jinja.")
        encontrados = []
        for clave in _TAG.findall(texto):
            if clave not in encontrados:
                encontrados.append(clave)
        return encontrados

    def importar_plantilla(self, source: str | Path, datos: dict, usuario: str = "") -> str:
        source = Path(source)
        claves = self.detectar_campos(source)
        if not claves:
            raise ValueError("El Word no contiene etiquetas con formato {{campo}}.")
        folder = get_word_templates_subdir("firmas")
        filename = self._safe_filename(datos.get("archivo_relativo") or source.name)
        target = folder / filename
        if source.resolve() != target.resolve():
            if target.exists() and self.sha256(target) != self.sha256(source):
                raise ValueError(f"Ya existe una plantilla distinta llamada {filename}.")
            if not target.exists():
                shutil.copy2(source, target)
        configurados = {str(c.get("clave")): dict(c) for c in datos.get("campos") or []}
        campos = []
        for orden, clave in enumerate(claves):
            campo = configurados.get(clave, {})
            campos.append({
                "clave": clave,
                "etiqueta": campo.get("etiqueta") or clave.replace("_", " ").title(),
                "origen": campo.get("origen") or "manual",
                "campo_origen": campo.get("campo_origen") or "",
                "tipo": campo.get("tipo") or "texto",
                "obligatorio": bool(campo.get("obligatorio")),
                "valor_defecto": campo.get("valor_defecto") or "",
                "orden": orden,
            })
        plantilla = dict(datos)
        plantilla.update({
            "id": datos.get("id") or str(uuid.uuid4()),
            "archivo_relativo": filename,
            "hash_docx": self.sha256(target),
            "campos": campos,
            "creado_por": usuario,
        })
        self._validar_configuracion(plantilla)
        return self.gestor.guardar_plantilla_firma(plantilla)

    def guardar_configuracion(self, plantilla: dict, usuario: str = "") -> str:
        actual = self.gestor.get_plantilla_firma(str(plantilla.get("id") or "")) or {}
        merged = {**actual, **dict(plantilla)}
        path = get_word_templates_subdir("firmas") / Path(str(merged.get("archivo_relativo") or "")).name
        claves = self.detectar_campos(path)
        campos = {str(c.get("clave")): dict(c) for c in merged.get("campos") or []}
        if set(claves) != set(campos):
            raise ValueError("Los campos configurados no coinciden con las etiquetas actuales del Word.")
        merged["hash_docx"] = self.sha256(path)
        merged["creado_por"] = usuario or merged.get("creado_por")
        self._validar_configuracion(merged)
        return self.gestor.guardar_plantilla_firma(merged)

    def comprobar_plantilla(self, plantilla: dict) -> None:
        path = self.ruta_plantilla(plantilla)
        if not path.is_file():
            raise FileNotFoundError(f"No se encuentra la plantilla compartida: {path}")
        actual_hash = self.sha256(path)
        if actual_hash != str(plantilla.get("hash_docx") or ""):
            raise ValueError(
                "La plantilla Word se ha modificado en red. Un administrador debe revisar sus campos y zonas."
            )
        claves = set(self.detectar_campos(path))
        configurados = {str(c.get("clave")) for c in plantilla.get("campos") or []}
        if claves != configurados:
            raise ValueError("Las etiquetas del Word ya no coinciden con la configuracion guardada.")

    def valores_iniciales(self, plantilla: dict, empresa: dict | None = None,
                          tercero: dict | None = None, usuario: str = "") -> dict:
        empresa = dict(empresa or {})
        tercero = dict(tercero or {})
        sistema = {
            "fecha_actual": date.today().strftime("%d/%m/%Y"),
            "usuario_nombre": usuario,
            "anio_actual": str(date.today().year),
        }
        fuentes = {"empresa": empresa, "tercero": tercero, "sistema": sistema}
        out = {}
        for campo in plantilla.get("campos") or []:
            origen = str(campo.get("origen") or "manual")
            key = str(campo.get("campo_origen") or campo.get("clave") or "")
            valor = campo.get("valor_defecto") or ""
            if origen in fuentes:
                fuente = fuentes[origen]
                clave_real = key
                if clave_real not in fuente:
                    clave_real = _ALIAS_CAMPO_ORIGEN.get(origen, {}).get(clave_real, clave_real)
                valor = fuente.get(clave_real, valor)
            elif origen == "fijo":
                valor = campo.get("valor_defecto") or ""
            out[str(campo["clave"])] = "" if valor is None else str(valor)
        return out

    def generar_documento(self, plantilla_id: str, valores: dict, *, empresa: dict | None = None,
                          tercero: dict | None = None, firmantes: list[dict] | None = None,
                          usuario: str = "", ejercicio: int | None = None) -> dict:
        plantilla = self.gestor.get_plantilla_firma(plantilla_id)
        if not plantilla or not bool(plantilla.get("activa")):
            raise ValueError("La plantilla no existe o no esta activa.")
        self.comprobar_plantilla(plantilla)
        valores = {str(k): v for k, v in dict(valores or {}).items()}
        self.validar_valores(plantilla, valores)
        documento_id = str(uuid.uuid4())
        codigo = str((empresa or {}).get("codigo") or (empresa or {}).get("codigo_empresa") or "")
        year = int(ejercicio or (empresa or {}).get("ejercicio") or date.today().year)
        output = self._directorio_generado(codigo, year, documento_id)
        titulo = self._safe_stem(plantilla.get("nombre") or "Documento")
        out_docx = output / f"{titulo}.docx"
        out_pdf = output / f"{titulo}.pdf"
        render_docx(str(self.ruta_plantilla(plantilla)), valores, str(out_docx))
        convert_docx_to_pdf(str(out_docx), str(out_pdf))
        documento = {
            "id": documento_id,
            "plantilla_id": plantilla["id"], "plantilla_version": plantilla["version"],
            "plantilla_hash": plantilla["hash_docx"], "codigo_empresa": codigo or None,
            "tercero_id": (tercero or {}).get("id") or (tercero or {}).get("tercero_id"),
            "titulo": plantilla.get("nombre") or titulo, "datos": valores,
            "firmantes": list(firmantes or []), "ruta_docx": str(out_docx),
            "ruta_pdf": str(out_pdf), "hash_docx": self.sha256(out_docx),
            "hash_pdf": self.sha256(out_pdf), "estado": "generado", "creado_por": usuario,
        }
        self.gestor.guardar_documento_firma_generado(documento)
        return documento

    def regenerar_pdf(self, documento_id: str) -> dict:
        documento = self.gestor.get_documento_firma_generado(documento_id)
        if not documento:
            raise ValueError("Documento generado no encontrado.")
        docx = Path(documento["ruta_docx"])
        pdf = Path(documento["ruta_pdf"])
        if not docx.is_file():
            raise FileNotFoundError("No se encuentra el Word generado.")
        convert_docx_to_pdf(str(docx), str(pdf))
        cambios = {"hash_docx": self.sha256(docx), "hash_pdf": self.sha256(pdf), "estado": "revisado"}
        self.gestor.actualizar_documento_firma_generado(documento_id, cambios)
        return {**documento, **cambios}

    @staticmethod
    def validar_valores(plantilla: dict, valores: dict) -> None:
        for campo in plantilla.get("campos") or []:
            clave = str(campo.get("clave") or "")
            valor = str(valores.get(clave) or "").strip()
            etiqueta = str(campo.get("etiqueta") or clave)
            if campo.get("obligatorio") and not valor:
                raise ValueError(f"El campo {etiqueta} es obligatorio.")
            if not valor:
                continue
            tipo = str(campo.get("tipo") or "texto")
            if tipo == "email" and not _EMAIL.match(valor):
                raise ValueError(f"El campo {etiqueta} no contiene un email valido.")
            if tipo == "nif" and not validar_nif_cif_nie(normalizar_nif_cif(valor)):
                raise ValueError(f"El campo {etiqueta} no contiene un NIF/CIF/NIE valido.")
            if tipo == "numero":
                try:
                    float(valor.replace(".", "").replace(",", "."))
                except ValueError as exc:
                    raise ValueError(f"El campo {etiqueta} debe ser numerico.") from exc
            if tipo == "fecha":
                formatos = ("%d/%m/%Y", "%Y-%m-%d")
                if not any(_fecha_valida(valor, formato) for formato in formatos):
                    raise ValueError(f"El campo {etiqueta} debe ser una fecha valida.")

    @staticmethod
    def ruta_plantilla(plantilla: dict) -> Path:
        return get_word_templates_subdir("firmas") / Path(str(plantilla.get("archivo_relativo") or "")).name

    @staticmethod
    def sha256(path: str | Path) -> str:
        digest = hashlib.sha256()
        with Path(path).open("rb") as fh:
            for chunk in iter(lambda: fh.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    @staticmethod
    def _safe_filename(value: str) -> str:
        name = _SAFE.sub("_", Path(str(value or "plantilla.docx")).name).strip(". ")
        if not name.lower().endswith(".docx"):
            name += ".docx"
        return name

    @staticmethod
    def _safe_stem(value: str) -> str:
        return _SAFE.sub("_", str(value or "Documento")).strip(". ") or "Documento"

    @classmethod
    def _directorio_generado(cls, codigo: str, ejercicio: int, documento_id: str) -> Path:
        root = get_document_repository_dir()
        if codigo:
            digits = "".join(ch for ch in codigo if ch.isdigit())
            root = root / "Empresas" / f"E{digits.zfill(5)[:5]}" / str(int(ejercicio))
        path = root / "Firmas" / "Generados" / documento_id
        path.mkdir(parents=True, exist_ok=False)
        return path

    @staticmethod
    def _validar_configuracion(plantilla: dict) -> None:
        if not str(plantilla.get("nombre") or "").strip():
            raise ValueError("La plantilla necesita un nombre.")
        if plantilla.get("alcance") not in {"global", "empresas"}:
            raise ValueError("El alcance de la plantilla no es valido.")
        if plantilla.get("alcance") == "empresas" and not plantilla.get("empresas"):
            raise ValueError("Selecciona al menos una empresa para el alcance restringido.")
        for campo in plantilla.get("campos") or []:
            if campo.get("origen") not in ORIGENES_CAMPO or campo.get("tipo") not in TIPOS_CAMPO:
                raise ValueError(f"Configuracion no valida para el campo {campo.get('clave')}.")
        ordenes = [int(f.get("orden") or 0) for f in plantilla.get("firmantes") or []]
        if ordenes and sorted(ordenes) != list(range(1, len(ordenes) + 1)):
            raise ValueError("El orden de los firmantes debe ser correlativo desde 1.")
        for firmante in plantilla.get("firmantes") or []:
            origen = str(firmante.get("origen") or "manual")
            email = str(firmante.get("email") or "").strip()
            if origen not in ORIGENES_FIRMANTE:
                raise ValueError(f"Origen no valido para el firmante {firmante.get('rol')}.")
            if origen == "manual" and not email:
                raise ValueError(f"El firmante manual {firmante.get('rol')} necesita un email.")
            if email and not _EMAIL.match(email):
                raise ValueError(f"El email configurado para {firmante.get('rol')} no es valido.")
        roles = {str(f.get("rol") or "") for f in plantilla.get("firmantes") or []}
        if any(str(z.get("rol") or "") not in roles for z in plantilla.get("zonas") or []):
            raise ValueError("Hay zonas asociadas a roles firmantes inexistentes.")


def _fecha_valida(value: str, formato: str) -> bool:
    try:
        datetime.strptime(value, formato)
        return True
    except ValueError:
        return False
