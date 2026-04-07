from __future__ import annotations

import os
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

from utils.utilidades import get_documentos_output_dir, get_documentos_output_structure, get_word_templates_dir


PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_.]+)\s*}}")


class DocumentosService:
    GLOBAL_CODE = "__global__"
    SAMPLE_TEMPLATES = (
        {
            "filename": "entrega_a_cuenta.docx",
            "nombre": "Entrega a cuenta",
            "tipo_documento": "entrega_a_cuenta",
            "descripcion": "Plantilla base para entregas a cuenta.",
        },
        {
            "filename": "contrato_prestamo_simple.docx",
            "nombre": "Contrato de prestamo simple",
            "tipo_documento": "contrato_prestamo_simple",
            "descripcion": "Plantilla base para prestamos simples.",
        },
        {
            "filename": "mandato_profesional_basico.docx",
            "nombre": "Mandato profesional basico",
            "tipo_documento": "mandato_profesional_basico",
            "descripcion": "Plantilla base de mandato profesional.",
        },
    )

    def __init__(self, gestor, codigo_empresa: str, ejercicio: int, global_mode: bool = False):
        self._gestor = gestor
        self._codigo = codigo_empresa
        self._ejercicio = ejercicio
        self._global_mode = bool(global_mode)
        self._empresa = self._load_empresa()

    def ensure_demo_templates(self) -> list[dict]:
        templates_dir = Path(self._templates_dir())
        templates_dir.mkdir(parents=True, exist_ok=True)
        created = []
        for item in self.SAMPLE_TEMPLATES:
            path = templates_dir / item["filename"]
            if not path.exists():
                self._build_sample_template(path, item["nombre"])
            variables = self.extract_variables_from_template(str(path))
            existing = self._gestor.buscar_plantilla_documento_por_nombre(self._codigo, self._ejercicio, item["nombre"])
            payload = {
                "id": existing.get("id") if existing else None,
                "codigo_empresa": self._codigo,
                "ejercicio": self._ejercicio,
                "nombre": item["nombre"],
                "tipo_documento": item["tipo_documento"],
                "descripcion": item["descripcion"],
                "ruta_template": str(path),
                "variables": variables,
                "activa": True,
            }
            plantilla_id = self._gestor.upsert_plantilla_documento(payload)
            created.append(self._gestor.get_plantilla_documento(plantilla_id))
        return created

    def register_template_file(self, file_path: str, *, nombre: str | None = None, tipo_documento: str | None = None):
        src = Path(file_path)
        if not src.exists():
            raise FileNotFoundError(f"No se encuentra la plantilla: {file_path}")
        dest_dir = Path(self._templates_dir())
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / src.name
        if src.resolve() != dest.resolve():
            shutil.copy2(src, dest)
        variables = self.extract_variables_from_template(str(dest))
        template_name = (nombre or src.stem.replace("_", " ").strip()).strip()
        existing = self._gestor.buscar_plantilla_documento_por_nombre(self._codigo, self._ejercicio, template_name)
        plantilla_id = self._gestor.upsert_plantilla_documento(
            {
                "id": existing.get("id") if existing else None,
                "codigo_empresa": self._codigo,
                "ejercicio": self._ejercicio,
                "nombre": template_name,
                "tipo_documento": tipo_documento or src.stem,
                "descripcion": "",
                "ruta_template": str(dest),
                "variables": variables,
                "activa": True,
            }
        )
        return self._gestor.get_plantilla_documento(plantilla_id)

    def refresh_template_variables(self, plantilla_id: int):
        plantilla = self._gestor.get_plantilla_documento(plantilla_id)
        if not plantilla:
            raise ValueError("Plantilla no encontrada.")
        plantilla["variables"] = self.extract_variables_from_template(plantilla.get("ruta_template"))
        self._gestor.upsert_plantilla_documento(plantilla)
        return self._gestor.get_plantilla_documento(plantilla_id)

    def rename_template(self, plantilla_id: int, nuevo_nombre: str):
        plantilla = self._gestor.get_plantilla_documento(plantilla_id)
        if not plantilla:
            raise ValueError("Plantilla no encontrada.")
        nombre = str(nuevo_nombre or "").strip()
        if not nombre:
            raise ValueError("El nombre de la plantilla es obligatorio.")
        plantilla["nombre"] = nombre
        self._gestor.upsert_plantilla_documento(plantilla)
        return self._gestor.get_plantilla_documento(plantilla_id)

    def extract_variables_from_template(self, template_path: str) -> list[str]:
        doc = Document(template_path)
        found = set()
        for text in self._iter_doc_text(doc):
            for match in PLACEHOLDER_RE.findall(text or ""):
                found.add(match.strip())
        return sorted(found)

    def generate_document(self, payload: dict) -> dict:
        plantilla = self._gestor.get_plantilla_documento(int(payload.get("plantilla_id")))
        if not plantilla:
            raise ValueError("Selecciona una plantilla documental.")

        cliente_data = dict(payload.get("cliente") or {})
        intervinientes = [dict(item) for item in (payload.get("intervinientes") or [])]
        custom_data = dict(payload.get("custom_data") or {})
        titulo = str(payload.get("titulo_documento") or plantilla.get("nombre") or "Documento").strip()
        if not titulo:
            raise ValueError("El titulo del documento es obligatorio.")

        now = datetime.now()
        context = self._build_context(
            titulo=titulo,
            cliente=cliente_data,
            intervinientes=intervinientes,
            operacion=payload.get("operacion") or {},
            custom_data=custom_data,
            fecha=now,
        )

        output_dir = self._build_output_dir(
            tipo_documento=plantilla.get("tipo_documento") or plantilla.get("nombre"),
            cliente_nombre=cliente_data.get("nombre_razon_social") or cliente_data.get("nombre") or "sin_cliente",
            operacion_titulo=(payload.get("operacion") or {}).get("titulo") or "",
        )
        output_dir.mkdir(parents=True, exist_ok=True)
        base_filename = self._safe_name(f"{now.strftime('%Y%m%d_%H%M%S')}_{titulo}") or f"documento_{now.strftime('%Y%m%d_%H%M%S')}"
        out_docx = output_dir / f"{base_filename}.docx"
        out_pdf = output_dir / f"{base_filename}.pdf"

        self.render_template_to_docx(plantilla.get("ruta_template"), context, str(out_docx))

        pdf_error = ""
        pdf_generated = False
        try:
            self._convert_docx_to_pdf(str(out_docx), str(out_pdf))
            pdf_generated = out_pdf.exists()
        except Exception as exc:
            pdf_error = str(exc)

        documento_id = self._gestor.upsert_documento_generado(
            {
                "codigo_empresa": self._codigo,
                "ejercicio": self._ejercicio,
                "plantilla_id": plantilla.get("id"),
                "cliente_id": cliente_data.get("cliente_id"),
                "operacion_id": payload.get("operacion_id"),
                "titulo_documento": titulo,
                "fecha_generacion": now.strftime("%Y-%m-%d %H:%M"),
                "ruta_docx": str(out_docx),
                "ruta_pdf": str(out_pdf) if pdf_generated else "",
                "estado": "generado" if pdf_generated else "docx_generado",
                "observaciones": payload.get("observaciones") or pdf_error,
                "json_datos_generacion": payload,
            }
        )

        linked = []
        has_cliente_in_list = any(str(item.get("_source") or "") == "cliente" for item in intervinientes)
        if not has_cliente_in_list:
            main_cliente = self._persist_interviniente_from_cliente(cliente_data)
            if main_cliente:
                linked.append({"interviniente_id": main_cliente, "rol_en_documento": "cliente"})
        for item in intervinientes:
            interviniente_id = self._persist_interviniente_snapshot(item)
            linked.append(
                {
                    "interviniente_id": interviniente_id,
                    "rol_en_documento": item.get("rol_en_documento") or item.get("rol") or "interviniente",
                }
            )
        self._gestor.set_documento_intervinientes(documento_id, linked)
        return self._gestor.get_documento_generado(documento_id)

    def build_preview_values(self, payload: dict) -> dict[str, str]:
        plantilla = None
        if payload.get("plantilla_id"):
            plantilla = self._gestor.get_plantilla_documento(int(payload.get("plantilla_id")))
        variables = list((plantilla or {}).get("variables") or [])
        context = self._build_context(
            titulo=str(payload.get("titulo_documento") or (plantilla or {}).get("nombre") or "Documento").strip(),
            cliente=dict(payload.get("cliente") or {}),
            intervinientes=[dict(item) for item in (payload.get("intervinientes") or [])],
            operacion=payload.get("operacion") or {},
            custom_data=dict(payload.get("custom_data") or {}),
            fecha=datetime.now(),
        )
        flat = self._flatten_context(context)
        out = {}
        for key in variables:
            out[key] = str(flat.get(key, payload.get("custom_data", {}).get(key, "")) or "")
        return out

    def duplicate_document_payload(self, documento_id: int) -> dict:
        doc = self._gestor.get_documento_generado(documento_id)
        if not doc:
            raise ValueError("Documento no encontrado.")
        payload = dict(doc.get("json_datos_generacion") or {})
        payload["titulo_documento"] = f"{payload.get('titulo_documento') or doc.get('titulo_documento') or 'Documento'} copia"
        return payload

    def render_template_to_docx(self, template_path: str, context: dict, output_path: str) -> None:
        doc = Document(template_path)
        flat = self._flatten_context(context)
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, flat)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, flat)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, flat)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, flat)
        doc.save(output_path)

    def build_interviniente_from_tercero(self, tercero: dict) -> dict:
        return {
            "tipo_persona": "juridica",
            "nombre_razon_social": tercero.get("nombre", ""),
            "nif": tercero.get("nif", ""),
            "domicilio": tercero.get("direccion", ""),
            "municipio": tercero.get("poblacion", ""),
            "provincia": tercero.get("provincia", ""),
            "cp": tercero.get("cp", ""),
            "telefono": tercero.get("telefono", ""),
            "email": tercero.get("email", ""),
            "representante": "",
            "cargo": "",
            "cliente_id": tercero.get("id"),
            "es_cliente_habitual": True,
            "observaciones": "",
        }

    def build_interviniente_from_empresa(self, empresa: dict) -> dict:
        return {
            "tipo_persona": "juridica",
            "nombre_razon_social": empresa.get("nombre", ""),
            "nif": empresa.get("cif", ""),
            "domicilio": empresa.get("direccion", ""),
            "municipio": empresa.get("poblacion", ""),
            "provincia": empresa.get("provincia", ""),
            "cp": empresa.get("cp", ""),
            "telefono": empresa.get("telefono", ""),
            "email": empresa.get("email", ""),
            "representante": "",
            "cargo": "",
            "cliente_id": f"empresa::{empresa.get('codigo','')}::{empresa.get('ejercicio','')}",
            "es_cliente_habitual": True,
            "observaciones": "",
        }

    def list_clientes(self) -> list[dict]:
        rows = []
        by_codigo = {}
        for empresa in self._gestor.listar_empresas():
            codigo = str(empresa.get("codigo") or "").strip()
            ejercicio = empresa.get("ejercicio")
            if not codigo or ejercicio is None:
                continue
            current = by_codigo.get(codigo)
            candidate = {
                "id": f"empresa::{codigo}::{ejercicio}",
                "codigo": codigo,
                "ejercicio": ejercicio,
                "nombre": empresa.get("nombre", ""),
                "nif": empresa.get("cif", ""),
                "poblacion": empresa.get("poblacion", ""),
                "direccion": empresa.get("direccion", ""),
                "cp": empresa.get("cp", ""),
                "provincia": empresa.get("provincia", ""),
                "telefono": empresa.get("telefono", ""),
                "email": empresa.get("email", ""),
            }
            if current is None:
                by_codigo[codigo] = candidate
                continue
            try:
                current_ej = int(current.get("ejercicio"))
            except Exception:
                current_ej = -1
            try:
                new_ej = int(ejercicio)
            except Exception:
                new_ej = -1
            if new_ej >= current_ej:
                by_codigo[codigo] = candidate
        rows = list(by_codigo.values())
        rows.sort(key=lambda item: (str(item.get("nombre") or "").lower(), str(item.get("codigo") or "")))
        return rows

    def list_contactos(self) -> list[dict]:
        contactos = []
        for item in self._gestor.listar_intervinientes(self._codigo, self._ejercicio, solo_habituales=True):
            row = dict(item)
            row["_contact_source"] = "interviniente"
            contactos.append(row)
        for tercero in self._gestor.listar_terceros():
            contactos.append(
                {
                    "id": f"tercero::{tercero.get('id')}",
                    "nombre_razon_social": tercero.get("nombre", ""),
                    "nif": tercero.get("nif", ""),
                    "email": tercero.get("email", ""),
                    "telefono": tercero.get("telefono", ""),
                    "municipio": tercero.get("poblacion", ""),
                    "provincia": tercero.get("provincia", ""),
                    "domicilio": tercero.get("direccion", ""),
                    "cp": tercero.get("cp", ""),
                    "cliente_id": tercero.get("id"),
                    "_contact_source": "tercero",
                }
            )
        contactos.sort(key=lambda item: str(item.get("nombre_razon_social") or "").lower())
        return contactos

    def _build_context(self, *, titulo: str, cliente: dict, intervinientes: list[dict], operacion: dict, custom_data: dict, fecha: datetime) -> dict:
        cliente_nombre = cliente.get("nombre_razon_social") or cliente.get("nombre") or ""
        lines = []
        indexed = {}
        for pos, item in enumerate(intervinientes, start=1):
            nombre = item.get("nombre_razon_social") or item.get("nombre") or ""
            rol = item.get("rol_en_documento") or item.get("rol") or "interviniente"
            if nombre:
                lines.append(f"{rol}: {nombre}")
            indexed[f"interviniente_{pos}"] = {
                "nombre_razon_social": nombre,
                "nif": item.get("nif", ""),
                "domicilio": item.get("domicilio", "") or item.get("direccion", ""),
                "municipio": item.get("municipio", "") or item.get("poblacion", ""),
                "provincia": item.get("provincia", ""),
                "cp": item.get("cp", ""),
                "telefono": item.get("telefono", ""),
                "email": item.get("email", ""),
                "representante": item.get("representante", ""),
                "cargo": item.get("cargo", ""),
                "rol_en_documento": rol,
            }
        context = {
            "fecha_hoy": fecha.strftime("%d/%m/%Y"),
            "empresa": {
                "nombre": self._empresa.get("nombre", ""),
                "cif": self._empresa.get("cif", ""),
                "direccion": self._empresa.get("direccion", ""),
                "cp": self._empresa.get("cp", ""),
                "poblacion": self._empresa.get("poblacion", ""),
                "provincia": self._empresa.get("provincia", ""),
                "telefono": self._empresa.get("telefono", ""),
                "email": self._empresa.get("email", ""),
            },
            "cliente": {
                "nombre_razon_social": cliente_nombre,
                "nif": cliente.get("nif", ""),
                "domicilio": cliente.get("domicilio", "") or cliente.get("direccion", ""),
                "municipio": cliente.get("municipio", "") or cliente.get("poblacion", ""),
                "provincia": cliente.get("provincia", ""),
                "cp": cliente.get("cp", ""),
                "telefono": cliente.get("telefono", ""),
                "email": cliente.get("email", ""),
                "representante": cliente.get("representante", ""),
                "cargo": cliente.get("cargo", ""),
            },
            "documento": {
                "titulo_documento": titulo,
                "observaciones": custom_data.get("observaciones", ""),
            },
            "operacion": {
                "titulo": operacion.get("titulo", ""),
                "tipo_operacion": operacion.get("tipo_operacion", ""),
                "descripcion": operacion.get("descripcion", ""),
            },
            "intervinientes_resumen": "\n".join(lines),
        }
        context.update(indexed)
        for key, value in custom_data.items():
            if "." not in str(key):
                context.setdefault("custom", {})[key] = value
        return context

    def _persist_interviniente_from_cliente(self, cliente_data: dict) -> int | None:
        nombre = cliente_data.get("nombre_razon_social") or cliente_data.get("nombre")
        if not nombre:
            return None
        if cliente_data.get("cliente_id"):
            for item in self._gestor.listar_intervinientes(self._codigo, self._ejercicio, solo_habituales=True):
                if str(item.get("cliente_id") or "") == str(cliente_data.get("cliente_id")):
                    return int(item.get("id"))
        payload = {
            "codigo_empresa": self._codigo,
            "ejercicio": self._ejercicio,
            "tipo_persona": cliente_data.get("tipo_persona") or "juridica",
            "nombre_razon_social": nombre,
            "nif": cliente_data.get("nif", ""),
            "domicilio": cliente_data.get("domicilio", "") or cliente_data.get("direccion", ""),
            "municipio": cliente_data.get("municipio", "") or cliente_data.get("poblacion", ""),
            "provincia": cliente_data.get("provincia", ""),
            "cp": cliente_data.get("cp", ""),
            "telefono": cliente_data.get("telefono", ""),
            "email": cliente_data.get("email", ""),
            "representante": cliente_data.get("representante", ""),
            "cargo": cliente_data.get("cargo", ""),
            "cliente_id": cliente_data.get("cliente_id"),
            "es_cliente_habitual": bool(cliente_data.get("cliente_id")),
            "observaciones": cliente_data.get("observaciones", ""),
        }
        return self._gestor.upsert_interviniente(payload)

    def _persist_interviniente_snapshot(self, interviniente: dict) -> int:
        payload = {
            "codigo_empresa": self._codigo,
            "ejercicio": self._ejercicio,
            "tipo_persona": interviniente.get("tipo_persona") or "fisica",
            "nombre_razon_social": interviniente.get("nombre_razon_social") or interviniente.get("nombre") or "",
            "nif": interviniente.get("nif", ""),
            "domicilio": interviniente.get("domicilio", "") or interviniente.get("direccion", ""),
            "municipio": interviniente.get("municipio", "") or interviniente.get("poblacion", ""),
            "provincia": interviniente.get("provincia", ""),
            "cp": interviniente.get("cp", ""),
            "telefono": interviniente.get("telefono", ""),
            "email": interviniente.get("email", ""),
            "representante": interviniente.get("representante", ""),
            "cargo": interviniente.get("cargo", ""),
            "cliente_id": interviniente.get("cliente_id"),
            "es_cliente_habitual": bool(interviniente.get("es_cliente_habitual")),
            "observaciones": interviniente.get("observaciones", ""),
        }
        return self._gestor.upsert_interviniente(payload)

    def _templates_dir(self) -> str:
        base_dir = Path(__file__).resolve().parents[1] / "plantillas"
        return get_word_templates_dir(str(base_dir))

    def _build_output_dir(self, *, tipo_documento: str, cliente_nombre: str, operacion_titulo: str) -> Path:
        base_dir = Path(__file__).resolve().parents[1] / "documentos_generados"
        root = Path(get_documentos_output_dir(str(base_dir)))
        structure = get_documentos_output_structure()
        empresa_dir = self._safe_name(self._empresa.get("nombre") or self._codigo) or self._codigo
        if structure == "operacion" and operacion_titulo:
            leaf = self._safe_name(operacion_titulo) or "sin_operacion"
        elif structure == "tipo_documento":
            leaf = self._safe_name(tipo_documento) or "documentos"
        else:
            leaf = self._safe_name(cliente_nombre) or "sin_cliente"
        return root / empresa_dir / leaf

    def _flatten_context(self, value: dict, prefix: str = "") -> dict[str, str]:
        out = {}
        for key, item in (value or {}).items():
            full_key = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(item, dict):
                out.update(self._flatten_context(item, full_key))
            else:
                out[full_key] = "" if item is None else str(item)
        return out

    def _replace_in_paragraph(self, paragraph, context: dict[str, str]) -> None:
        text = paragraph.text or ""
        if "{{" not in text:
            return
        updated = text
        for match in PLACEHOLDER_RE.findall(text):
            updated = updated.replace(f"{{{{ {match} }}}}", context.get(match, ""))
            updated = updated.replace(f"{{{{{match}}}}}", context.get(match, ""))
        if updated != text:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated

    def _iter_doc_text(self, doc: Document):
        for paragraph in doc.paragraphs:
            yield paragraph.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph.text
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph.text
            for paragraph in section.footer.paragraphs:
                yield paragraph.text

    def _build_sample_template(self, path: Path, titulo: str) -> None:
        doc = Document()
        doc.add_heading(titulo, level=1)
        doc.add_paragraph("Fecha: {{ fecha_hoy }}")
        doc.add_paragraph("Empresa: {{ empresa.nombre }} - {{ empresa.cif }}")
        doc.add_paragraph("Cliente: {{ cliente.nombre_razon_social }}")
        doc.add_paragraph("NIF: {{ cliente.nif }}")
        doc.add_paragraph("Domicilio: {{ cliente.domicilio }}, {{ cliente.cp }} {{ cliente.municipio }} ({{ cliente.provincia }})")
        doc.add_paragraph("Documento: {{ documento.titulo_documento }}")
        doc.add_paragraph("Operacion: {{ operacion.titulo }}")
        doc.add_paragraph("Intervinientes adicionales:")
        doc.add_paragraph("{{ intervinientes_resumen }}")
        doc.add_paragraph("Interviniente 1: {{ interviniente_1.rol_en_documento }} - {{ interviniente_1.nombre_razon_social }}")
        doc.add_paragraph("Interviniente 2: {{ interviniente_2.rol_en_documento }} - {{ interviniente_2.nombre_razon_social }}")
        doc.add_paragraph("Observaciones: {{ documento.observaciones }}")
        doc.save(path)

    def _safe_name(self, value: str) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        for ch in '<>:"/\\|?*':
            text = text.replace(ch, " ")
        return " ".join(text.split())

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        from docx2pdf import convert

        convert(docx_path, pdf_path)

    def _load_empresa(self) -> dict:
        if self._global_mode:
            return {
                "codigo": self.GLOBAL_CODE,
                "nombre": "Documentos globales",
                "cif": "",
                "direccion": "",
                "cp": "",
                "poblacion": "",
                "provincia": "",
                "telefono": "",
                "email": "",
            }
        try:
            return self._gestor.get_empresa(self._codigo, self._ejercicio) or {}
        except Exception:
            return {}
