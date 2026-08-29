from __future__ import annotations

import hashlib
import json
import os
import re
import secrets
import shutil
import uuid
import webbrowser
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZIP_DEFLATED, ZipFile

from services.tramites_dgt_repository import DgtRepository
from services.tramites_dgt_documentos import MIME_POR_EXTENSION_DGT
from utils.utilidades import get_app_data_dir, get_word_templates_subdir
from utils.validaciones import normalizar_nif_cif, validar_nif_cif_nie


TIPO_COMPRAVENTA = "compraventa_cambio_titularidad"
DOCUMENTOS_BASE = (
    ("contrato_compraventa", "Contrato de compraventa"),
    ("mandato_dgt_comprador", "Mandato DGT comprador"),
)
ROLES_PARTE = {"vendedor", "comprador"}
TEMPLATE_FILENAMES = {
    "contrato_compraventa": "dgt_contrato_compraventa.docx",
    "mandato_dgt_comprador": "dgt_mandato_comprador.docx",
}


def get_protocol_url_from_argv(argv: list[str]) -> str:
    for arg in argv[1:] or []:
        raw = str(arg or "").strip()
        if raw.lower().startswith("gest2a3eco://tramites-dgt/"):
            return raw
    return ""


@dataclass(slots=True)
class LinkSeguro:
    rol: str
    token: str
    url: str


class TramitesDgtService:
    def __init__(
        self,
        gestor=None,
        session=None,
        repository: DgtRepository | None = None,
        firma_client=None,
        firma_gestor_email: str = "",
        firma_gestor_telefono: str = "",
        almacenamiento_client=None,
        almacenamiento_base_path: str = "",
    ):
        if repository is None:
            raise ValueError(
                "TramitesDgtService necesita el backend de integraciones configurado."
            )
        self._repo = repository
        self._session = session
        self._firma_client = firma_client
        self._firma_gestor_email = str(
            firma_gestor_email or getattr(firma_client, "from_email", "") or ""
        ).strip()
        self._firma_gestor_telefono = str(firma_gestor_telefono or "").strip()
        self._almacenamiento_client = almacenamiento_client
        self._almacenamiento_base_path = str(almacenamiento_base_path or "").strip().rstrip("/")

    def crear_expediente_minimo(self, payload: dict) -> str:
        create_remote = getattr(self._repo, "create_expediente", None)
        if callable(create_remote):
            return create_remote(payload)
        expediente_id = str(uuid.uuid4())
        vendedor = self._crear_link("vendedor")
        comprador = self._crear_link("comprador")
        referencia = self._siguiente_referencia()
        expediente = {
            "id": expediente_id,
            "referencia": referencia,
            "tipo": TIPO_COMPRAVENTA,
            "estado": "borrador",
            "titulo": str(payload.get("titulo") or "").strip() or referencia,
            "vendedor_nombre": str(payload.get("vendedor_nombre") or "").strip(),
            "vendedor_email": str(payload.get("vendedor_email") or "").strip(),
            "vendedor_telefono": self._normalizar_telefono(payload.get("vendedor_telefono")),
            "comprador_nombre": str(payload.get("comprador_nombre") or "").strip(),
            "comprador_email": str(payload.get("comprador_email") or "").strip(),
            "comprador_telefono": self._normalizar_telefono(payload.get("comprador_telefono")),
            "vehiculo_matricula": self._normalizar_matricula(payload.get("vehiculo_matricula")),
            "vehiculo_bastidor": str(payload.get("vehiculo_bastidor") or "").strip().upper(),
            "precio_venta": self._parse_float(payload.get("precio_venta")),
            "fecha_operacion": str(payload.get("fecha_operacion") or "").strip(),
            "observaciones": str(payload.get("observaciones") or "").strip(),
            "vendedor_token_hash": self._hash_token(vendedor.token),
            "comprador_token_hash": self._hash_token(comprador.token),
            "vendedor_token_created_at": self._now(),
            "comprador_token_created_at": self._now(),
            "firma_estado": "pendiente",
            "created_by": getattr(getattr(self._session, "user", None), "id", None),
        }
        self._repo.upsert_expediente(expediente)
        return expediente_id

    def guardar_expediente(self, expediente_id: str, payload: dict) -> None:
        actual = self._repo.get_expediente(expediente_id)
        if not actual:
            raise ValueError("Expediente DGT no encontrado.")
        self._asegurar_editable(actual)
        codigo_tasa = str(payload.get("codigo_tasa") or "").strip().upper()
        if codigo_tasa and not re.fullmatch(r"[A-Z0-9-]{6,32}", codigo_tasa):
            raise ValueError("El codigo de tasa debe tener entre 6 y 32 letras, numeros o guiones.")
        actualizado = dict(actual)
        for key in (
            "titulo",
            "vendedor_nombre",
            "vendedor_email",
            "comprador_nombre",
            "comprador_email",
            "vehiculo_bastidor",
            "fecha_operacion",
            "observaciones",
        ):
            actualizado[key] = str(payload.get(key) or "").strip()
        actualizado["codigo_tasa"] = codigo_tasa
        actualizado["modelo_620_presentado"] = bool(payload.get("modelo_620_presentado"))
        actualizado["vendedor_telefono"] = self._normalizar_telefono(payload.get("vendedor_telefono"))
        actualizado["comprador_telefono"] = self._normalizar_telefono(payload.get("comprador_telefono"))
        actualizado["vehiculo_matricula"] = self._normalizar_matricula(payload.get("vehiculo_matricula"))
        actualizado["precio_venta"] = self._parse_float(payload.get("precio_venta"))
        if actualizado.get("estado") == "validado":
            actualizado["estado"] = "revision"
            actualizado["validado_por"] = None
            actualizado["validado_at"] = None
        self._repo.upsert_expediente(actualizado)

    def listar_expedientes(self) -> list[dict]:
        return self._repo.listar_expedientes()

    def get_expediente(self, expediente_id: str) -> dict | None:
        return self._repo.get_expediente(expediente_id)

    def get_links(self, expediente: dict) -> dict[str, str]:
        ref = expediente.get("referencia") or expediente.get("id")
        return {
            "vendedor": self._build_url("vendedor", ref),
            "comprador": self._build_url("comprador", ref),
        }

    def regenerar_links(self, expediente_id: str) -> dict[str, str]:
        create_remote = getattr(self._repo, "create_links", None)
        if callable(create_remote):
            return create_remote(expediente_id)
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        vendedor = self._crear_link("vendedor")
        comprador = self._crear_link("comprador")
        expediente["vendedor_token_hash"] = self._hash_token(vendedor.token)
        expediente["comprador_token_hash"] = self._hash_token(comprador.token)
        expediente["vendedor_token_created_at"] = self._now()
        expediente["comprador_token_created_at"] = self._now()
        self._repo.upsert_expediente(expediente)
        ref = expediente.get("referencia") or expediente_id
        return {
            "vendedor": self._build_url("vendedor", ref, vendedor.token),
            "comprador": self._build_url("comprador", ref, comprador.token),
        }

    def revocar_link(self, expediente_id: str, rol: str) -> None:
        rol = self._validar_rol(rol)
        revoke_remote = getattr(self._repo, "revoke_link", None)
        if callable(revoke_remote):
            revoke_remote(expediente_id, rol)
            return
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        expediente[f"{rol}_token_hash"] = ""
        self._repo.upsert_expediente(expediente)

    def solicitar_subsanacion(self, expediente_id: str, rol: str, mensaje: str) -> dict:
        rol = self._validar_rol(rol)
        mensaje = str(mensaje or "").strip()
        if len(mensaje) < 3:
            raise ValueError("Indica el motivo de la subsanacion.")
        request_remote = getattr(self._repo, "solicitar_subsanacion", None)
        if callable(request_remote):
            return request_remote(expediente_id, rol, mensaje)
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        link = self._crear_link(rol)
        expediente[f"{rol}_token_hash"] = self._hash_token(link.token)
        expediente[f"{rol}_token_created_at"] = self._now()
        expediente["estado"] = "requiere_subsanacion"
        expediente[f"{rol}_estado"] = "requiere_subsanacion"
        expediente.setdefault("subsanaciones", []).append(
            {"rol": rol, "mensaje": mensaje, "created_at": self._now()}
        )
        self._repo.upsert_expediente(expediente)
        ref = expediente.get("referencia") or expediente_id
        return {"status": "pendiente", "rol": rol, "url": self._build_url(rol, ref, link.token)}

    def verificar_token(self, referencia: str, rol: str, token: str) -> dict:
        rol = self._validar_rol(rol)
        expediente = self._repo.get_expediente_por_referencia(referencia)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        stored_hash = str(expediente.get(f"{rol}_token_hash") or "")
        if not stored_hash or not secrets.compare_digest(stored_hash, self._hash_token(token)):
            raise PermissionError("Enlace DGT no valido o caducado.")
        return expediente

    def completar_desde_link(self, referencia: str, rol: str, token: str, payload: dict) -> None:
        expediente = self.verificar_token(referencia, rol, token)
        self.guardar_datos_parte(expediente["id"], rol, payload)

    def parse_link_seguro(self, url: str) -> dict:
        parsed = urlparse(str(url or "").strip())
        if parsed.scheme != "gest2a3eco" or parsed.netloc != "tramites-dgt":
            raise ValueError("El enlace no corresponde a Trámites DGT.")
        parts = [unquote(part) for part in parsed.path.strip("/").split("/") if part]
        if len(parts) < 2:
            raise ValueError("El enlace DGT no contiene rol y referencia.")
        rol = self._validar_rol(parts[0])
        referencia = parts[1]
        token = (parse_qs(parsed.query).get("token") or [""])[0]
        if not token:
            raise ValueError("El enlace DGT no contiene token. Regenera el enlace antes de abrir el formulario.")
        return {"rol": rol, "referencia": referencia, "token": token}

    def guardar_datos_parte(self, expediente_id: str, rol: str, payload: dict) -> None:
        rol = self._validar_rol(rol)
        actual = self._repo.get_expediente(expediente_id)
        if not actual:
            raise ValueError("Expediente DGT no encontrado.")
        self._asegurar_editable(actual)
        update_remote = getattr(self._repo, "update_parte", None)
        if callable(update_remote):
            update_remote(expediente_id, rol, self._normalizar_payload_parte(payload))
            return
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        datos = self._normalizar_payload_parte(payload)
        expediente[f"{rol}_payload"] = datos
        if datos.get("nombre"):
            expediente[f"{rol}_nombre"] = datos["nombre"]
        if datos.get("email"):
            expediente[f"{rol}_email"] = datos["email"]
        if datos.get("telefono"):
            expediente[f"{rol}_telefono"] = datos["telefono"]
        if rol == "vendedor":
            if datos.get("vehiculo_matricula"):
                expediente["vehiculo_matricula"] = datos["vehiculo_matricula"]
            if datos.get("vehiculo_bastidor"):
                expediente["vehiculo_bastidor"] = datos["vehiculo_bastidor"]
            if datos.get("precio_venta") is not None:
                expediente["precio_venta"] = datos["precio_venta"]
            if datos.get("fecha_operacion"):
                expediente["fecha_operacion"] = datos["fecha_operacion"]
        if expediente.get("estado") in ("borrador", ""):
            expediente["estado"] = "pendiente_revision"
        if expediente.get("estado") == "validado":
            expediente["estado"] = "revision"
            expediente["validado_por"] = None
            expediente["validado_at"] = None
        self._repo.upsert_expediente(expediente)

    def adjuntar_documento(self, expediente_id: str, rol: str, file_path: str, tipo: str = "", descripcion: str = "") -> dict:
        rol = str(rol or "").strip().lower()
        if rol not in ROLES_PARTE | {"gestor"}:
            raise ValueError("Rol DGT no valido.")
        path = Path(file_path).expanduser()
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"No existe el documento adjunto: {path}")
        if path.suffix.lower() not in MIME_POR_EXTENSION_DGT:
            raise ValueError("Solo se admiten archivos PDF, JPG, JPEG y PNG.")
        tipo = str(tipo or "").strip() or "documentacion"
        if len(tipo) > 64:
            raise ValueError("El tipo de documento no puede superar 64 caracteres.")
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        upload_remote = getattr(self._repo, "upload_documento", None)
        if callable(upload_remote):
            return upload_remote(expediente_id, rol, tipo, str(path))
        digest = self._hash_file(path)
        item = {
            "id": str(uuid.uuid4()),
            "rol": rol,
            "tipo": tipo,
            "descripcion": str(descripcion or "").strip(),
            "nombre_archivo": path.name,
            "ruta": str(path.resolve()),
            "sha256": digest,
            "added_at": self._now(),
        }
        documentos = list(expediente.get("documentos") or [])
        documentos.append(item)
        expediente["documentos"] = documentos
        if tipo == "modelo_620":
            expediente["modelo_620_presentado"] = True
        if expediente.get("estado") == "validado":
            expediente["estado"] = "revision"
            expediente["validado_por"] = None
            expediente["validado_at"] = None
        self._repo.upsert_expediente(expediente)
        return item

    def validar_expediente(self, expediente_id: str) -> None:
        if getattr(self._repo, "online", False):
            user_id = getattr(getattr(self._session, "user", None), "id", 0) or 0
            self._repo.validar_expediente(expediente_id, user_id)
            return
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        errors = self.validar_datos(expediente)
        if errors:
            raise ValueError("\n".join(errors))
        user_id = getattr(getattr(self._session, "user", None), "id", 0) or 0
        self._repo.validar_expediente(expediente_id, user_id)

    def validar_datos(self, expediente: dict) -> list[str]:
        errors = []
        if not str(expediente.get("vehiculo_matricula") or "").strip():
            errors.append("La matricula del vehiculo es obligatoria.")
        if not str(expediente.get("vendedor_nombre") or "").strip():
            errors.append("El nombre del vendedor es obligatorio.")
        if not str(expediente.get("comprador_nombre") or "").strip():
            errors.append("El nombre del comprador es obligatorio.")
        for rol in ("vendedor", "comprador"):
            payload = expediente.get(f"{rol}_payload") or {}
            nif = normalizar_nif_cif(payload.get("nif") or payload.get("dni") or "")
            if not nif:
                errors.append(f"El NIF/NIE/CIF del {rol} es obligatorio.")
            elif not validar_nif_cif_nie(nif):
                errors.append(f"El NIF/NIE/CIF del {rol} no es valido.")
            for campo, etiqueta in (
                ("nombre", "nombre"),
                ("email", "correo electronico"),
                ("telefono", "telefono"),
                ("direccion", "direccion"),
                ("cp", "codigo postal"),
                ("poblacion", "poblacion"),
                ("provincia", "provincia"),
            ):
                if not str(payload.get(campo) or "").strip():
                    errors.append(f"El {etiqueta} del {rol} es obligatorio.")
            if payload.get("email") and not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", str(payload["email"])):
                errors.append(f"El correo electronico del {rol} no es valido.")
            if payload.get("telefono") and not re.fullmatch(r"\d{9,15}", str(payload["telefono"])):
                errors.append(f"El telefono del {rol} debe contener entre 9 y 15 digitos.")
            if payload.get("cp") and not re.fullmatch(r"\d{5}", str(payload["cp"])):
                errors.append(f"El codigo postal del {rol} debe tener 5 digitos.")
            if payload.get("tipo_persona") == "juridica":
                if not payload.get("representante_nombre"):
                    errors.append(f"El nombre del representante del {rol} es obligatorio.")
                representante_nif = normalizar_nif_cif(payload.get("representante_nif") or "")
                if not representante_nif:
                    errors.append(f"El DNI/NIE del representante del {rol} es obligatorio.")
                elif not validar_nif_cif_nie(representante_nif):
                    errors.append(f"El DNI/NIE del representante del {rol} no es valido.")
        vendedor = expediente.get("vendedor_payload") or {}
        for campo, etiqueta in (
            ("vehiculo_matricula", "matricula"),
            ("vehiculo_bastidor", "bastidor"),
            ("vehiculo_marca", "marca"),
            ("vehiculo_modelo", "modelo"),
            ("primera_matriculacion", "fecha de primera matriculacion"),
            ("kilometraje", "kilometraje"),
            ("precio_venta", "precio de venta"),
            ("fecha_operacion", "fecha de entrega"),
            ("hora_entrega", "hora de entrega"),
            ("forma_pago", "forma de pago"),
            ("llaves_vehiculo", "numero de llaves"),
        ):
            if vendedor.get(campo) in (None, ""):
                errors.append(f"La {etiqueta} del vehiculo es obligatoria.")
        bastidor = str(vendedor.get("vehiculo_bastidor") or "")
        if bastidor and not re.fullmatch(r"[A-HJ-NPR-Z0-9]{17}", bastidor):
            errors.append("El bastidor debe contener 17 caracteres validos.")
        comprador = expediente.get("comprador_payload") or {}
        for campo, etiqueta in (
            ("direccion_envio", "direccion"),
            ("cp_envio", "codigo postal"),
            ("poblacion_envio", "poblacion"),
            ("provincia_envio", "provincia"),
        ):
            if not str(comprador.get(campo) or "").strip():
                errors.append(f"La {etiqueta} de envio es obligatoria.")
        if comprador.get("cp_envio") and not re.fullmatch(r"\d{5}", str(comprador["cp_envio"])):
            errors.append("El codigo postal de envio debe tener 5 digitos.")
        if expediente.get("modelo_620_presentado") and not any(
            doc.get("tipo") == "modelo_620" for doc in expediente.get("documentos") or []
        ):
            errors.append("Debes adjuntar el modelo 620 presentado.")
        if vendedor.get("tipo_persona") == "juridica":
            tiene_factura = any(
                doc.get("rol") == "vendedor" and doc.get("tipo") == "factura"
                for doc in expediente.get("documentos") or []
            )
            if not tiene_factura:
                errors.append("La factura es obligatoria cuando el vendedor es persona juridica.")
        return errors

    def generar_documentos(self, expediente_id: str) -> list[dict]:
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        if expediente.get("estado") != "validado":
            raise ValueError("Valida expresamente el expediente antes de generar documentos.")
        out = []
        documentos = list(DOCUMENTOS_BASE)
        vendedor = expediente.get("vendedor_payload") or {}
        if vendedor.get("tipo_persona") == "juridica":
            documentos = [item for item in documentos if item[0] != "contrato_compraventa"]
        fecha_generacion = self._now()
        version = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        for tipo, titulo in documentos:
            generated = self._generar_documento(expediente, tipo, titulo, version)
            contexto = self._document_context(expediente)
            remotos = self._guardar_documentos_remotos(
                expediente, "Generados", generated
            )
            if remotos:
                contexto["dataprius"] = remotos
            doc_id = self._repo.insertar_documento_generado(
                {
                    "expediente_id": expediente_id,
                    "tipo_documento": tipo,
                    "titulo": titulo,
                    "fecha_generacion": fecha_generacion,
                    "ruta_docx": generated.get("ruta_docx"),
                    "ruta_pdf": generated.get("ruta_pdf"),
                    "ruta_txt": generated.get("ruta_txt"),
                    "json_datos_generacion": contexto,
                    "hash_contenido": generated.get("hash_contenido"),
                    "estado": generated.get("estado"),
                }
            )
            out.append(
                {
                    "id": doc_id,
                    "tipo_documento": tipo,
                    "titulo": titulo,
                    "fecha_generacion": fecha_generacion,
                    **generated,
                }
            )
        return out

    def preparar_paquete_firma(self, expediente_id: str, provider: str = "") -> dict:
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        documentos = self._repo.listar_documentos_generados(expediente_id)
        rutas = []
        tipos_incluidos = set()
        for doc in documentos:
            tipo = str(doc.get("tipo_documento") or "")
            if tipo in tipos_incluidos:
                continue
            ruta = doc.get("ruta_pdf") or doc.get("ruta_docx") or doc.get("ruta_txt")
            if ruta:
                rutas.append(ruta)
                tipos_incluidos.add(tipo)
        if not rutas:
            raise ValueError("Genera los documentos antes de preparar la firma.")
        expediente["firma_estado"] = "preparado"
        expediente["firma_provider"] = str(provider or "").strip()
        expediente["firma_request_id"] = ""
        self._repo.upsert_expediente(expediente)
        return {"expediente_id": expediente_id, "provider": expediente["firma_provider"], "documentos": rutas}

    def enviar_a_firma(self, expediente_id: str, usar_sms: bool = False) -> dict:
        if self._firma_client is None:
            raise ValueError(
                "Configura signrequest_token y signrequest_from_email para enviar documentos."
            )
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        if expediente.get("firma_estado") in {"enviado", "parcialmente_firmado"}:
            raise ValueError("Ya existe una solicitud de firma activa para este expediente.")
        documentos = self._ultimos_documentos_por_tipo(expediente_id)
        if not documentos:
            raise ValueError("Genera los documentos antes de enviarlos a firma.")
        vendedor = expediente.get("vendedor_payload") or {}
        comprador = expediente.get("comprador_payload") or {}
        solicitudes = []
        for doc in documentos:
            tipo = str(doc.get("tipo_documento") or "")
            ruta = doc.get("ruta_pdf") or doc.get("ruta_docx") or doc.get("ruta_txt")
            firmantes = [self._firmante(comprador, 1)]
            if tipo == "contrato_compraventa":
                if str(vendedor.get("email") or "").strip().lower() == str(
                    comprador.get("email") or ""
                ).strip().lower():
                    raise ValueError(
                        "Vendedor y comprador deben tener emails distintos para firmar el contrato."
                    )
                firmantes = [self._firmante(vendedor, 1), self._firmante(comprador, 2)]
            elif tipo == "mandato_dgt_comprador":
                if not self._firma_gestor_email:
                    raise ValueError(
                        "Configura signrequest_gestor_email para que el mandatario firme el mandato."
                    )
                if self._firma_gestor_email.lower() == str(
                    comprador.get("email") or ""
                ).strip().lower():
                    raise ValueError(
                        "Comprador y mandatario deben tener emails distintos para firmar el mandato."
                    )
                # SignRequest reserva el indice de etiqueta 0 al remitente.
                # Lo enviamos primero en la lista de contactos, aunque firme
                # despues del comprador mediante el orden de firma.
                firmantes = [
                    {
                        "email": self._firma_gestor_email,
                        "telefono": self._firma_gestor_telefono,
                        "order": 2,
                    },
                    self._firmante(comprador, 1),
                ]
            resultado = self._firma_client.enviar_documento(
                ruta=ruta,
                firmantes=firmantes,
                asunto=f"{doc.get('titulo') or tipo} - {expediente.get('referencia')}",
                mensaje=(
                    f"Firma del documento {doc.get('titulo') or tipo} "
                    f"del expediente {expediente.get('referencia')}."
                ),
                external_id=f"{expediente_id}:{doc.get('id')}",
                usar_sms=usar_sms,
            )
            if not resultado.get("uuid"):
                raise RuntimeError("SignRequest no devolvio el identificador de la solicitud.")
            solicitudes.append(
                {
                    "request_id": resultado["uuid"],
                    "documento_id": str(doc.get("id")),
                    "tipo_documento": tipo,
                    "estado": resultado.get("status") or "sent",
                }
            )
        expediente["firma_estado"] = "enviado"
        expediente["firma_provider"] = "signrequest"
        expediente["firma_request_id"] = json.dumps(
            [item["request_id"] for item in solicitudes], ensure_ascii=True
        )
        expediente["firma_evidencia"] = {
            "request_id": expediente["firma_request_id"],
            "solicitudes": solicitudes,
            "enviado_at": self._now(),
        }
        self._repo.upsert_expediente(expediente)
        return {"estado": "enviado", "solicitudes": solicitudes}

    def actualizar_estado_firma(self, expediente_id: str) -> dict:
        if self._firma_client is None:
            raise ValueError("SignRequest no esta configurado.")
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        evidencia = dict(expediente.get("firma_evidencia") or {})
        solicitudes = list(evidencia.get("solicitudes") or [])
        if not solicitudes:
            raise ValueError("El expediente no tiene solicitudes de firma.")
        estados = []
        avisos = []
        for solicitud in solicitudes:
            data = self._firma_client.consultar(solicitud["request_id"])
            estado = str(data.get("status") or "desconocido")
            solicitud["estado"] = estado
            solicitud["actualizado_at"] = self._now()
            if estado.lower() in {"signed", "completed"} and not solicitud.get("ruta_firmado"):
                tipo = str(solicitud.get("tipo_documento") or "documento")
                evidencias = self._firma_client.descargar_evidencias(
                    solicitud["request_id"],
                    str(self._output_dir(expediente) / "firmados"),
                    f"{tipo}_{solicitud['request_id'][:8]}",
                )
                solicitud.update(evidencias)
                avisos.extend(self._registrar_evidencias_firma(expediente_id, tipo, evidencias))
            estados.append(estado)
        normalizados = {estado.lower() for estado in estados}
        if normalizados and normalizados <= {"signed", "completed"}:
            estado_global = "firmado"
        elif normalizados & {"declined", "cancelled", "expired"}:
            estado_global = "incidencia"
        elif normalizados & {"signed", "completed"}:
            estado_global = "parcialmente_firmado"
        else:
            estado_global = "enviado"
        evidencia["solicitudes"] = solicitudes
        evidencia["actualizado_at"] = self._now()
        expediente["firma_estado"] = estado_global
        expediente["firma_provider"] = "signrequest"
        expediente["firma_evidencia"] = evidencia
        self._repo.upsert_expediente(expediente)
        return {"estado": estado_global, "solicitudes": solicitudes, "avisos": avisos}

    def finalizar_expediente(self, expediente_id: str) -> dict:
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        if expediente.get("estado") == "terminado":
            return expediente
        if str(expediente.get("firma_estado") or "").lower() not in {
            "firmado", "signed", "completed"
        }:
            raise ValueError("Solo se puede terminar un expediente con la firma completada.")
        finalizar_remoto = getattr(self._repo, "finalizar_expediente", None)
        if callable(finalizar_remoto):
            return finalizar_remoto(expediente_id)
        expediente["estado"] = "terminado"
        self._repo.upsert_expediente(expediente)
        return expediente

    @staticmethod
    def _asegurar_editable(expediente: dict) -> None:
        if str(expediente.get("estado") or "").lower() == "terminado":
            raise ValueError("El expediente esta terminado y es de solo lectura.")

    def anular_ultima_firma(self, expediente_id: str) -> dict:
        if self._firma_client is None:
            raise ValueError("SignRequest no esta configurado.")
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        evidencia = dict(expediente.get("firma_evidencia") or {})
        solicitudes = list(evidencia.get("solicitudes") or [])
        if not solicitudes:
            raise ValueError("El expediente no tiene una solicitud de firma activa.")

        estados = []
        for solicitud in solicitudes:
            data = self._firma_client.consultar(solicitud["request_id"])
            estado = str(data.get("status") or "desconocido").lower()
            solicitud["estado"] = estado
            estados.append(estado)
        if any(estado in {"signed", "completed"} for estado in estados):
            raise ValueError("No se puede anular porque alguno de los documentos ya esta firmado.")

        anuladas = 0
        for solicitud in solicitudes:
            if solicitud["estado"] not in {"cancelled", "canceled", "declined", "expired"}:
                self._firma_client.cancelar(solicitud["request_id"])
                anuladas += 1
            solicitud["estado"] = "cancelled"
            solicitud["cancelado_at"] = self._now()
        evidencia["solicitudes"] = solicitudes
        evidencia["cancelado_at"] = self._now()
        expediente["firma_estado"] = "cancelado"
        expediente["firma_request_id"] = ""
        expediente["firma_evidencia"] = evidencia
        self._repo.upsert_expediente(expediente)
        return {"estado": "cancelado", "solicitudes_anuladas": anuladas}

    def _ultimos_documentos_por_tipo(self, expediente_id: str) -> list[dict]:
        seleccionados = []
        tipos = set()
        tipos_firmables = {tipo for tipo, _titulo in DOCUMENTOS_BASE}
        for doc in self._repo.listar_documentos_generados(expediente_id):
            tipo = str(doc.get("tipo_documento") or "")
            if tipo in tipos_firmables and tipo not in tipos:
                seleccionados.append(doc)
                tipos.add(tipo)
        return seleccionados

    def _registrar_evidencias_firma(self, expediente_id: str, tipo: str, evidencias: dict) -> list[str]:
        expediente = self._repo.get_expediente(expediente_id) or {}
        existentes = self._repo.listar_documentos_generados(expediente_id)
        fecha = self._now()
        nuevos = (
            (
                f"{tipo}_firmado",
                f"{tipo.replace('_', ' ').title()} firmado",
                evidencias.get("ruta_firmado"),
                evidencias.get("sha256_firmado"),
            ),
            (
                f"{tipo}_registro_firma",
                f"Registro de firma - {tipo.replace('_', ' ')}",
                evidencias.get("ruta_registro_firma"),
                evidencias.get("sha256_registro_firma"),
            ),
        )
        rutas_existentes = {
            str(doc.get("ruta_pdf") or "") for doc in existentes if doc.get("ruta_pdf")
        }
        avisos = []
        for tipo_documento, titulo, ruta, sha256 in nuevos:
            if not ruta or str(ruta) in rutas_existentes:
                continue
            remotos = []
            error_dataprius = ""
            try:
                remotos = self._guardar_documentos_remotos(
                    expediente, "Firmados", {"ruta_pdf": ruta}
                )
            except Exception as exc:
                error_dataprius = str(exc)
                avisos.append(
                    f"{Path(str(ruta)).name} se guardo localmente, pero no en Dataprius: {exc}"
                )
            self._repo.insertar_documento_generado(
                {
                    "expediente_id": expediente_id,
                    "tipo_documento": tipo_documento,
                    "titulo": titulo,
                    "fecha_generacion": fecha,
                    "ruta_pdf": str(ruta),
                    "ruta_docx": None,
                    "ruta_txt": None,
                    "json_datos_generacion": {
                        "origen": "signrequest",
                        "dataprius": remotos,
                        "dataprius_error": error_dataprius,
                    },
                    "hash_contenido": sha256,
                    "estado": "firmado",
                }
            )
        return avisos

    def _guardar_documentos_remotos(
        self, expediente: dict, categoria: str, documentos: dict
    ) -> list[dict]:
        if self._almacenamiento_client is None or not self._almacenamiento_base_path:
            return []
        referencia = str(expediente.get("referencia") or expediente.get("id") or "").strip()
        ruta_remota = f"{self._almacenamiento_base_path}/{referencia}/{categoria}"
        rutas = []
        for campo in ("ruta_docx", "ruta_pdf", "ruta_txt"):
            ruta = str(documentos.get(campo) or "").strip()
            if ruta and ruta not in rutas:
                rutas.append(ruta)
        return [
            self._almacenamiento_client.subir_archivo(ruta_remota, ruta)
            for ruta in rutas
        ]

    @staticmethod
    def _firmante(datos: dict, order: int) -> dict:
        email = str(datos.get("email") or "").strip()
        if not email:
            raise ValueError("Todos los firmantes deben tener un email guardado.")
        return {
            "email": email,
            "telefono": str(datos.get("telefono") or "").strip(),
            "order": order,
        }

    def listar_documentos(self, expediente_id: str) -> list[dict]:
        return self._repo.listar_documentos_generados(expediente_id)

    def eliminar_expediente(self, expediente_id: str) -> None:
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        generados = self._repo.listar_documentos_generados(expediente_id)
        self._repo.eliminar_expediente(expediente_id)
        for doc in generados:
            self._eliminar_ficheros_generados(doc)
        output_dir = self._output_dir_path(expediente)
        if output_dir.is_dir():
            try:
                output_dir.rmdir()
            except OSError:
                pass

    def eliminar_documento_generado(self, documento_id) -> None:
        doc = self._repo.eliminar_documento_generado(documento_id)
        if not doc:
            raise ValueError("Documento generado no encontrado.")
        self._eliminar_ficheros_generados(doc)

    def eliminar_documento_aportado(self, expediente_id: str, documento_id: str) -> None:
        expediente = self._repo.get_expediente(expediente_id)
        if not expediente:
            raise ValueError("Expediente DGT no encontrado.")
        delete_remote = getattr(self._repo, "eliminar_documento_aportado", None)
        if callable(delete_remote):
            delete_remote(documento_id)
            return
        documentos = list(expediente.get("documentos") or [])
        restantes = [doc for doc in documentos if str(doc.get("id")) != str(documento_id)]
        if len(restantes) == len(documentos):
            raise ValueError("Documento aportado no encontrado.")
        expediente["documentos"] = restantes
        self._repo.upsert_expediente(expediente)

    def descargar_documento_aportado(self, documento_id: str, target_path: str) -> str:
        download = getattr(self._repo, "download_documento", None)
        if not callable(download):
            raise ValueError("La descarga bajo demanda solo esta disponible en el repositorio online.")
        return download(documento_id, target_path)

    def get_templates_dir(self) -> Path:
        return get_word_templates_subdir("tramites_dgt")

    def listar_plantillas_editables(self) -> list[dict]:
        base = self.get_templates_dir()
        out = []
        for tipo, titulo in DOCUMENTOS_BASE:
            filename = TEMPLATE_FILENAMES[tipo]
            path = base / filename
            out.append(
                {
                    "tipo_documento": tipo,
                    "titulo": titulo,
                    "filename": filename,
                    "path": str(path),
                    "exists": path.exists(),
                }
            )
        return out

    def ensure_plantillas_editables(self, overwrite: bool = False) -> list[dict]:
        base = self.get_templates_dir()
        packaged = Path(__file__).resolve().parents[1] / "plantillas" / "tramites_dgt"
        created = []
        for tipo, titulo in DOCUMENTOS_BASE:
            path = base / TEMPLATE_FILENAMES[tipo]
            if path.exists() and not overwrite:
                continue
            source = packaged / TEMPLATE_FILENAMES[tipo]
            if source.exists() and source.resolve() != path.resolve():
                path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(source, path)
                ok = path.exists()
            elif source.exists():
                ok = True
            else:
                ok = self._create_editable_template(tipo, titulo, path)
            created.append(
                {
                    "tipo_documento": tipo,
                    "titulo": titulo,
                    "path": str(path),
                    "created": ok,
                }
            )
        return created

    def abrir_carpeta_plantillas(self) -> Path:
        path = self.get_templates_dir()
        self._open_path(path)
        return path

    def abrir_plantilla(self, tipo: str) -> Path:
        info = {item["tipo_documento"]: item for item in self.listar_plantillas_editables()}
        item = info.get(str(tipo or "").strip())
        if not item:
            raise ValueError("Tipo de plantilla DGT no valido.")
        path = Path(item["path"])
        if not path.exists():
            self._create_editable_template(item["tipo_documento"], item["titulo"], path)
        self._open_path(path)
        return path

    def abrir_whatsapp(self, telefono: str, mensaje: str) -> None:
        phone = self._normalizar_telefono(telefono)
        if not phone:
            raise ValueError("No hay telefono para WhatsApp.")
        webbrowser.open(f"https://wa.me/{phone}?text={quote(mensaje)}")

    def _crear_link(self, rol: str) -> LinkSeguro:
        token = secrets.token_urlsafe(32)
        return LinkSeguro(rol=rol, token=token, url="")

    def _build_url(self, rol: str, referencia: str, token: str = "") -> str:
        base = f"gest2a3eco://tramites-dgt/{rol}/{quote(str(referencia))}"
        return f"{base}?token={quote(token)}" if token else base

    def _hash_token(self, token: str) -> str:
        return hashlib.sha256(str(token).encode("utf-8")).hexdigest()

    def _siguiente_referencia(self) -> str:
        prefix = f"DGT-{datetime.now().year}-"
        rows = self._repo.listar_expedientes()
        nums = []
        for row in rows:
            ref = str(row.get("referencia") or "")
            if ref.startswith(prefix):
                try:
                    nums.append(int(ref.rsplit("-", 1)[1]))
                except Exception:
                    pass
        return f"{prefix}{(max(nums) + 1) if nums else 1:04d}"

    def _output_dir(self, expediente: dict) -> Path:
        path = self._output_dir_path(expediente)
        path.mkdir(parents=True, exist_ok=True)
        return path

    def _output_dir_path(self, expediente: dict) -> Path:
        ref = str(expediente.get("referencia") or expediente.get("id") or "sin_ref").replace("/", "_")
        return get_app_data_dir() / "tramites_dgt" / ref

    def _eliminar_ficheros_generados(self, doc: dict) -> None:
        root = (get_app_data_dir() / "tramites_dgt").resolve()
        for key in ("ruta_docx", "ruta_pdf", "ruta_txt"):
            raw = doc.get(key)
            if not raw:
                continue
            path = Path(raw).resolve()
            if root not in path.parents:
                continue
            if path.is_file():
                path.unlink()

    def _generar_documento(self, expediente: dict, tipo: str, titulo: str, version: str) -> dict:
        output_dir = self._output_dir(expediente)
        nombre_versionado = f"{tipo}_{version}"
        txt_path = output_dir / f"{nombre_versionado}.txt"
        contenido = self._render_documento_texto(expediente, titulo)
        txt_path.write_text(contenido, encoding="utf-8")
        result = {
            "ruta_txt": str(txt_path),
            "ruta_docx": None,
            "ruta_pdf": None,
            "hash_contenido": hashlib.sha256(contenido.encode("utf-8")).hexdigest(),
            "estado": "generado_txt",
        }
        docx_path = output_dir / f"{nombre_versionado}.docx"
        if self._render_docx(expediente, tipo, titulo, docx_path):
            result["ruta_docx"] = str(docx_path)
            result["hash_contenido"] = self._hash_file(docx_path)
            result["estado"] = "generado_docx"
            pdf_path = output_dir / f"{nombre_versionado}.pdf"
            if self._convertir_pdf(docx_path, pdf_path):
                result["ruta_pdf"] = str(pdf_path)
                result["hash_contenido"] = self._hash_file(pdf_path)
                result["estado"] = "generado_pdf"
        return result

    def _render_docx(self, expediente: dict, tipo: str, titulo: str, out_docx_path: Path) -> bool:
        template_path = self._buscar_template(tipo)
        if not template_path:
            raise FileNotFoundError(f"No se encuentra la plantilla oficial para {titulo}.")
        context = self._document_context(expediente)
        context["titulo_documento"] = titulo
        from procesos.facturas_word import render_docx

        render_docx(str(template_path), context, str(out_docx_path))
        self._asegurar_etiquetas_firma(out_docx_path, tipo)
        return out_docx_path.exists()

    @staticmethod
    def _asegurar_etiquetas_firma(docx_path: Path, tipo: str) -> None:
        """Inserta campos invisibles que SignRequest convierte en zonas de firma."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        document = Document(str(docx_path))

        def add_tag(paragraph, signer_index: int) -> None:
            marker = f"[[s|{signer_index}"
            if marker in paragraph.text:
                return
            run = paragraph.add_run(f"[[s|{signer_index}{' ' * 26}]]")
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(18)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if tipo == "contrato_compraventa":
            if not document.tables or len(document.tables[-1].rows[-1].cells) < 2:
                raise ValueError(
                    "La plantilla del contrato necesita una tabla final con "
                    "las zonas de firma de vendedor y comprador."
                )
            cells = document.tables[-1].rows[-1].cells
            for cell, signer_index in zip(cells[:2], (1, 2)):
                if not any(
                    f"[[s|{signer_index}" in paragraph.text
                    for paragraph in cell.paragraphs
                ):
                    add_tag(cell.add_paragraph(), signer_index)
        elif tipo == "mandato_dgt_comprador":
            assignments = {"EL MANDANTE": 1, "EL MANDATARIO": 0}
            for heading, signer_index in assignments.items():
                if any(
                    f"[[s|{signer_index}" in paragraph.text
                    for paragraph in document.paragraphs
                ):
                    continue
                heading_index = next(
                    (
                        index
                        for index, paragraph in enumerate(document.paragraphs)
                        if paragraph.text.strip().upper() == heading
                    ),
                    None,
                )
                if heading_index is None:
                    raise ValueError(
                        f"La plantilla del mandato no contiene la zona '{heading}'."
                    )
                target = next(
                    (
                        paragraph
                        for paragraph in document.paragraphs[heading_index + 1 :]
                        if not paragraph.text.strip()
                    ),
                    None,
                )
                if target is None:
                    raise ValueError(
                        f"La plantilla del mandato necesita espacio de firma tras '{heading}'."
                    )
                add_tag(target, signer_index)
        document.save(str(docx_path))

    def _buscar_template(self, tipo: str) -> Path | None:
        filename = TEMPLATE_FILENAMES.get(tipo)
        if not filename:
            return None
        path = self.get_templates_dir() / filename
        return path if path.exists() and path.is_file() else None

    def _render_basic_docx(self, context: dict, titulo: str, out_docx_path: Path) -> bool:
        try:
            from docx import Document
        except Exception:
            return self._render_minimal_docx(context, titulo, out_docx_path)
        try:
            doc = Document()
            doc.add_heading(titulo, level=1)
            rows = self._document_rows(context)
            table = doc.add_table(rows=0, cols=2)
            for label, value in rows:
                cells = table.add_row().cells
                cells[0].text = str(label)
                cells[1].text = "" if value is None else str(value)
            doc.add_paragraph("")
            doc.add_paragraph(
                "Documento generado por Gest2A3Eco. Sustituir por plantilla oficial antes de firma electronica."
            )
            out_docx_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out_docx_path))
            return out_docx_path.exists()
        except Exception:
            return self._render_minimal_docx(context, titulo, out_docx_path)

    def _create_editable_template(self, tipo: str, titulo: str, path: Path) -> bool:
        context = self._empty_template_context()
        context["titulo_documento"] = titulo
        return self._render_basic_docx(context, titulo, path)

    def _empty_template_context(self) -> dict:
        return {
            "id": "",
            "referencia": "{{ referencia }}",
            "estado": "{{ estado }}",
            "titulo": "{{ titulo }}",
            "titulo_documento": "{{ titulo_documento }}",
            "vendedor_nombre": "{{ vendedor_nombre }}",
            "vendedor_nif": "{{ vendedor_nif }}",
            "comprador_nombre": "{{ comprador_nombre }}",
            "comprador_nif": "{{ comprador_nif }}",
            "vehiculo_matricula": "{{ vehiculo_matricula }}",
            "vehiculo_bastidor": "{{ vehiculo_bastidor }}",
            "precio_venta": "{{ precio_venta }}",
            "fecha_operacion": "{{ fecha_operacion }}",
            "observaciones": "{{ observaciones }}",
            "documentos_count": "{{ documentos_count }}",
            "documentos": [],
        }

    def _document_rows(self, context: dict) -> list[tuple[str, object]]:
        return [
            ("Referencia", context.get("referencia")),
            ("Vendedor", context.get("vendedor_nombre")),
            ("NIF vendedor", context.get("vendedor_nif")),
            ("Comprador", context.get("comprador_nombre")),
            ("NIF comprador", context.get("comprador_nif")),
            ("Matricula", context.get("vehiculo_matricula")),
            ("Bastidor", context.get("vehiculo_bastidor")),
            ("Precio venta", context.get("precio_venta")),
            ("Fecha operacion", context.get("fecha_operacion")),
            ("Documentos adjuntos", context.get("documentos_count")),
        ]

    def _render_minimal_docx(self, context: dict, titulo: str, out_docx_path: Path) -> bool:
        paragraphs = [titulo]
        paragraphs.extend(f"{label}: {'' if value is None else value}" for label, value in self._document_rows(context))
        paragraphs.append("Documento generado por Gest2A3Eco. Sustituir por plantilla oficial antes de firma electronica.")
        body = "".join(f"<w:p><w:r><w:t>{xml_escape(str(text))}</w:t></w:r></w:p>" for text in paragraphs)
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{body}<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/>"
            '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr></w:body>'
            "</w:document>"
        )
        content_types = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/>'
            "</Relationships>"
        )
        try:
            out_docx_path.parent.mkdir(parents=True, exist_ok=True)
            with ZipFile(out_docx_path, "w", ZIP_DEFLATED) as zf:
                zf.writestr("[Content_Types].xml", content_types)
                zf.writestr("_rels/.rels", rels)
                zf.writestr("word/document.xml", document_xml)
            return out_docx_path.exists()
        except Exception:
            return False

    def _convertir_pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        try:
            from procesos.facturas_word import convert_docx_to_pdf

            convert_docx_to_pdf(str(docx_path), str(pdf_path))
            return pdf_path.exists()
        except Exception:
            return False

    def _render_documento_texto(self, expediente: dict, titulo: str) -> str:
        ctx = self._document_context(expediente)
        return (
            f"{titulo}\n"
            f"Referencia: {ctx['referencia']}\n"
            f"Estado: {ctx['estado']}\n\n"
            f"Vendedor: {ctx['vendedor_nombre']}\n"
            f"NIF vendedor: {ctx['vendedor_nif']}\n"
            f"Comprador: {ctx['comprador_nombre']}\n"
            f"NIF comprador: {ctx['comprador_nif']}\n"
            f"Vehiculo: {ctx['vehiculo_matricula']} / {ctx['vehiculo_bastidor']}\n"
            f"Precio: {ctx['precio_venta']}\n"
            f"Fecha operacion: {ctx['fecha_operacion']}\n\n"
            f"Documentos adjuntos: {ctx['documentos_count']}\n\n"
            "Documento preliminar generado por Gest2A3Eco. Pendiente de plantilla DOCX/PDF y firma electronica.\n"
        )

    def _document_context(self, expediente: dict) -> dict:
        keys = (
            "id",
            "referencia",
            "estado",
            "titulo",
            "vendedor_nombre",
            "comprador_nombre",
            "vehiculo_matricula",
            "vehiculo_bastidor",
            "precio_venta",
            "fecha_operacion",
            "observaciones",
        )
        ctx = {key: expediente.get(key) for key in keys}
        vendedor_payload = expediente.get("vendedor_payload") or {}
        comprador_payload = expediente.get("comprador_payload") or {}
        self._add_parte_context(ctx, "vendedor", vendedor_payload)
        self._add_parte_context(ctx, "comprador", comprador_payload)
        ctx["vendedor_comparecencia"] = self._comparecencia(vendedor_payload, "vendedor")
        ctx["comprador_comparecencia"] = self._comparecencia(comprador_payload, "comprador")
        ctx["comprador_mandante"] = self._mandante(comprador_payload)
        ctx["vendedor_firma"] = self._firma_parte(vendedor_payload)
        ctx["comprador_firma"] = self._firma_parte(comprador_payload)
        ctx["fecha"] = self._fecha_es(expediente.get("fecha_operacion"))
        ctx["fecha_firma"] = ctx["fecha"] or datetime.now().strftime("%d/%m/%Y")
        ctx["marca_vehiculo"] = vendedor_payload.get("vehiculo_marca") or ""
        ctx["modelo_vehiculo"] = vendedor_payload.get("vehiculo_modelo") or ""
        ctx["matricula_vehiculo"] = vendedor_payload.get("vehiculo_matricula") or ctx.get("vehiculo_matricula") or ""
        ctx["bastidor_vehiculo"] = vendedor_payload.get("vehiculo_bastidor") or ctx.get("vehiculo_bastidor") or ""
        ctx["primeramatriculacion_vehiculo"] = self._fecha_es(
            vendedor_payload.get("primera_matriculacion") or vendedor_payload.get("vehiculo_primera_matriculacion")
        )
        ctx["kilometros_vehiculo"] = (
            vendedor_payload.get("kilometraje") or vendedor_payload.get("vehiculo_kilometros") or ""
        )
        ctx["precio_compra"] = self._importe_es(vendedor_payload.get("precio_venta") or ctx.get("precio_venta"))
        ctx["forma_pago"] = vendedor_payload.get("forma_pago") or ""
        ctx["llaves_vehiculo"] = vendedor_payload.get("llaves_vehiculo") or vendedor_payload.get("numero_llaves") or ""
        ctx["documentos"] = expediente.get("documentos") or []
        ctx["documentos_count"] = len(ctx["documentos"])
        return ctx

    def _add_parte_context(self, ctx: dict, rol: str, payload: dict) -> None:
        for campo in (
            "tipo_persona", "nombre", "nif", "email", "telefono", "direccion",
            "cp", "poblacion", "provincia", "representante_nombre", "representante_nif",
        ):
            ctx[f"{rol}_{campo}"] = payload.get(campo) or ""
        ctx[f"{rol}_direccion_completa"] = self._direccion_completa(payload)
        if rol == "comprador":
            envio = {
                "direccion": payload.get("direccion_envio") or payload.get("envio_direccion"),
                "cp": payload.get("cp_envio") or payload.get("envio_cp"),
                "poblacion": payload.get("poblacion_envio") or payload.get("envio_poblacion"),
                "provincia": payload.get("provincia_envio") or payload.get("envio_provincia"),
            }
            ctx["comprador_direccion_envio"] = self._direccion_completa(envio)

    def _comparecencia(self, payload: dict, rol: str) -> str:
        contacto = (
            f"con domicilio en {self._direccion_completa(payload)}, telefono "
            f"{payload.get('telefono') or ''} y correo electronico {payload.get('email') or ''}"
        )
        if payload.get("tipo_persona") == "juridica":
            return (
                f"{payload.get('nombre') or ''}, con CIF "
                f"{payload.get('nif') or ''}, {contacto}, representada en este acto por "
                f"Don/Doña {payload.get('representante_nombre') or ''}, con DNI/NIE "
                f"{payload.get('representante_nif') or ''}"
            )
        return (
            f"Don/Doña {payload.get('nombre') or ''}, con DNI/NIE {payload.get('nif') or ''}, "
            f"{contacto}"
        )

    def _mandante(self, payload: dict) -> str:
        domicilio = self._direccion_completa(
            {
                "direccion": payload.get("direccion_envio") or payload.get("envio_direccion") or payload.get("direccion"),
                "cp": payload.get("cp_envio") or payload.get("envio_cp") or payload.get("cp"),
                "poblacion": payload.get("poblacion_envio") or payload.get("envio_poblacion") or payload.get("poblacion"),
                "provincia": payload.get("provincia_envio") or payload.get("envio_provincia") or payload.get("provincia"),
            }
        )
        if payload.get("tipo_persona") == "juridica":
            return (
                f"La entidad {payload.get('nombre') or ''}, con CIF {payload.get('nif') or ''}, "
                f"con domicilio para notificaciones en {domicilio}, representada en este acto por "
                f"D./Dña. {payload.get('representante_nombre') or ''}, con DNI/NIE "
                f"{payload.get('representante_nif') or ''}, actuando en calidad de MANDANTE,"
            )
        return (
            f"Yo, D./Dña. {payload.get('nombre') or ''}, con DNI/NIE {payload.get('nif') or ''}, "
            f"con domicilio para notificaciones en {domicilio}, actuando en calidad de MANDANTE,"
        )

    def _firma_parte(self, payload: dict) -> str:
        if payload.get("tipo_persona") == "juridica":
            empresa = str(payload.get("nombre") or "").strip()
            representante = str(payload.get("representante_nombre") or "").strip()
            return "\n".join(
                parte for parte in (empresa, f"D./Dña. {representante}" if representante else "")
                if parte
            )
        return str(payload.get("nombre") or "")

    def _direccion_completa(self, payload: dict) -> str:
        localidad = " ".join(
            part for part in (
                str(payload.get("cp") or "").strip(),
                str(payload.get("poblacion") or "").strip(),
            ) if part
        )
        return ", ".join(
            part for part in (
                str(payload.get("direccion") or "").strip(),
                localidad,
                str(payload.get("provincia") or "").strip(),
            ) if part
        )

    def _fecha_es(self, value) -> str:
        raw = str(value or "").strip()
        try:
            return datetime.strptime(raw[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            return raw

    def _importe_es(self, value) -> str:
        try:
            return f"{float(value):,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
        except (TypeError, ValueError):
            return str(value or "")

    def _parse_float(self, value):
        raw = str(value or "").strip().replace(" ", "").replace("€", "")
        if not raw:
            return None
        if "," in raw and "." in raw:
            if raw.rfind(",") > raw.rfind("."):
                raw = raw.replace(".", "").replace(",", ".")
            else:
                raw = raw.replace(",", "")
        elif "," in raw:
            raw = raw.replace(".", "").replace(",", ".")
        elif "." in raw:
            partes = raw.split(".")
            if len(partes) > 2:
                raw = "".join(partes) if len(partes[-1]) == 3 else "".join(partes[:-1]) + "." + partes[-1]
            elif len(partes[-1]) == 3 and partes[0].lstrip("+-").isdigit():
                raw = "".join(partes)
        try:
            return float(raw)
        except Exception:
            return None

    def _normalizar_matricula(self, value) -> str:
        return "".join(ch for ch in str(value or "").upper() if ch.isalnum())

    def _normalizar_telefono(self, value) -> str:
        raw = "".join(ch for ch in str(value or "") if ch.isdigit() or ch == "+")
        if raw.startswith("+"):
            return raw[1:]
        return raw

    def _now(self) -> str:
        return datetime.utcnow().replace(microsecond=0).isoformat()

    def _validar_rol(self, rol: str) -> str:
        rol = str(rol or "").strip().lower()
        if rol not in ROLES_PARTE:
            raise ValueError("Rol DGT no valido.")
        return rol

    def _normalizar_payload_parte(self, payload: dict) -> dict:
        out = {
            "tipo_persona": str(payload.get("tipo_persona") or "fisica").strip().lower(),
            "nombre": str(payload.get("nombre") or "").strip(),
            "nif": normalizar_nif_cif(payload.get("nif") or payload.get("dni") or ""),
            "email": str(payload.get("email") or "").strip(),
            "telefono": self._normalizar_telefono(payload.get("telefono")),
            "direccion": str(payload.get("direccion") or "").strip(),
            "cp": str(payload.get("cp") or "").strip(),
            "poblacion": str(payload.get("poblacion") or "").strip(),
            "provincia": str(payload.get("provincia") or "").strip(),
            "representante_nombre": str(
                payload.get("representante_nombre") or payload.get("representante") or ""
            ).strip(),
            "representante_nif": normalizar_nif_cif(payload.get("representante_nif") or ""),
            "observaciones": str(payload.get("observaciones") or "").strip(),
        }
        out["vehiculo_matricula"] = self._normalizar_matricula(payload.get("vehiculo_matricula"))
        out["vehiculo_bastidor"] = str(payload.get("vehiculo_bastidor") or "").strip().upper()
        out["precio_venta"] = self._parse_float(payload.get("precio_venta"))
        aliases = {
            "vehiculo_marca": ("vehiculo_marca",),
            "vehiculo_modelo": ("vehiculo_modelo",),
            "primera_matriculacion": ("primera_matriculacion", "vehiculo_primera_matriculacion"),
            "kilometraje": ("kilometraje", "vehiculo_kilometros"),
            "fecha_operacion": ("fecha_operacion",),
            "hora_entrega": ("hora_entrega",),
            "forma_pago": ("forma_pago",),
            "llaves_vehiculo": ("llaves_vehiculo", "numero_llaves"),
            "direccion_envio": ("direccion_envio", "envio_direccion"),
            "cp_envio": ("cp_envio", "envio_cp"),
            "poblacion_envio": ("poblacion_envio", "envio_poblacion"),
            "provincia_envio": ("provincia_envio", "envio_provincia"),
        }
        for campo, opciones in aliases.items():
            out[campo] = str(next((payload.get(key) for key in opciones if payload.get(key) not in (None, "")), "")).strip()
        return out

    def _hash_file(self, path: Path) -> str:
        h = hashlib.sha256()
        with path.open("rb") as fh:
            for chunk in iter(lambda: fh.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()

    def _open_path(self, path: Path) -> None:
        try:
            os.startfile(str(path))  # type: ignore[attr-defined]
        except AttributeError:
            webbrowser.open(path.resolve().as_uri())
